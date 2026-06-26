import csv
import html
import json
import math
import re
import shutil
import time
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

import requests

try:
    from docx import Document
except ImportError:  # pragma: no cover - validation reports this if missing.
    Document = None


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
DOCS_DIR = ROOT / "docs"
PAPER_DIR = ROOT / "paper"
CACHE_DIR = DATA_DIR / "cache"

START_YEAR = 1900
END_YEAR = 2026
YEARS = list(range(START_YEAR, END_YEAR + 1))
YEAR_RANGE_TEXT = f"{START_YEAR}-{END_YEAR}"
YEAR_FILE_STEM = f"{START_YEAR}_{END_YEAR}"

TARGET_PER_YEAR = 100
CANDIDATES_PER_YEAR = 1000
GENERATED_DATE = date.today().isoformat()

S2_BULK_URL = "https://api.semanticscholar.org/graph/v1/paper/search/bulk"
S2_FIELDS = ",".join(
    [
        "paperId",
        "title",
        "year",
        "authors",
        "venue",
        "publicationVenue",
        "publicationDate",
        "citationCount",
        "influentialCitationCount",
        "abstract",
        "url",
        "externalIds",
        "openAccessPdf",
        "s2FieldsOfStudy",
        "publicationTypes",
    ]
)
S2_QUERIES = [
    "brain",
    "neuroscience brain",
    "cerebral cortex",
    "neural brain",
    "human brain",
    "mouse brain",
    "brain imaging",
    "brain connectivity",
    "cerebral",
    "cerebrum",
    "cortex",
    "cortical",
    "cerebellum",
    "medulla oblongata",
    "brain stem",
    "brainstem",
    "forebrain",
    "midbrain",
    "hindbrain",
    "diencephalon",
    "thalamus",
    "hypothalamus",
    "basal ganglia",
    "spinal cord brain",
    "central nervous system",
    "neurology brain",
    "neuron brain",
    "neurones brain",
    "neuroglia",
    "nerve cells brain",
    "cerebrospinal",
    "encephalon",
    "encephalitis brain",
    "aphasia brain",
    "visual cortex",
    "motor cortex",
    "sensory cortex",
]
S2_MAX_PAGES_PER_QUERY = 1

PAPERS_JSON = f"papers_{YEAR_FILE_STEM}.json"
PAPERS_CSV = f"papers_{YEAR_FILE_STEM}.csv"
CANDIDATES_JSON = f"candidates_top{CANDIDATES_PER_YEAR}_{YEAR_FILE_STEM}.json"
CANDIDATES_CSV = f"candidates_top{CANDIDATES_PER_YEAR}_{YEAR_FILE_STEM}.csv"
TAXONOMY_CSV = f"papers_taxonomy_{YEAR_FILE_STEM}.csv"
PERIOD_ANALYSIS_JSON = f"period_analysis_{YEAR_FILE_STEM}.json"
OVERALL_ANALYSIS_JSON = f"overall_analysis_{YEAR_FILE_STEM}.json"
GITHUB_LINKS_JSON = f"github_links_{YEAR_FILE_STEM}.json"
LINK_AUDIT_JSON = f"link_audit_{YEAR_FILE_STEM}.json"
SKILL2_PROVENANCE_JSON = "paper_curation_skill2_provenance.json"

LANGUAGES = {
    "en": "English",
    "ko": "한국어",
    "zh": "中文",
    "ja": "日本語",
}

UI_LABELS = {
    "en": {
        "papers": "papers",
        "years": "active years",
        "citations": "citations",
        "categories": "categories",
        "keyIdea": "Key idea",
        "strengths": "Strengths",
        "limitations": "Limitations",
        "noKeyword": "No keyword selected.",
    },
    "ko": {
        "papers": "논문",
        "years": "활성 연도",
        "citations": "인용",
        "categories": "분류",
        "keyIdea": "핵심 아이디어",
        "strengths": "장점",
        "limitations": "한계",
        "noKeyword": "선택된 키워드가 없습니다.",
    },
    "zh": {
        "papers": "论文",
        "years": "活跃年份",
        "citations": "引用",
        "categories": "分类",
        "keyIdea": "核心思想",
        "strengths": "优势",
        "limitations": "局限性",
        "noKeyword": "未选择关键词。",
    },
    "ja": {
        "papers": "論文",
        "years": "対象年",
        "citations": "引用",
        "categories": "分類",
        "keyIdea": "主要アイデア",
        "strengths": "強み",
        "limitations": "限界",
        "noKeyword": "キーワード未選択です。",
    },
}

RELEVANCE_TERMS = [
    "brain",
    "cerebral",
    "cerebrum",
    "cortex",
    "cortical",
    "neuron",
    "neurone",
    "neural",
    "neuroscience",
    "neurology",
    "neuroglia",
    "neuroimaging",
    "hippocampus",
    "amygdala",
    "cerebellum",
    "medulla oblongata",
    "brain stem",
    "brainstem",
    "forebrain",
    "midbrain",
    "hindbrain",
    "diencephalon",
    "thalamus",
    "hypothalamus",
    "basal ganglia",
    "central nervous system",
    "cerebrospinal",
    "encephalon",
    "encephalitis",
    "prefrontal",
    "connectome",
    "synapse",
    "glia",
    "microglia",
    "eeg",
    "meg",
    "ecog",
    "fmri",
    "mri",
]

CATEGORIES = [
    {
        "name": "Neuroimaging and Brain Mapping",
        "slug": "neuroimaging-and-brain-mapping",
        "patterns": [
            "mri",
            "fmri",
            "magnetic resonance",
            "neuroimaging",
            "brain mapping",
            "diffusion tensor",
            "tractography",
            "pet",
            "connectome",
            "functional connectivity",
        ],
        "accent": "#2563eb",
        "secondary": "#14b8a6",
        "overview": [
            "High-citation brain imaging work maps structure, function, connectivity, and disease signatures across MRI, fMRI, PET, DTI, and multimodal cohorts.",
            "The strongest recent trend is large-scale population neuroimaging with harmonized preprocessing, shared atlases, and open datasets.",
            "Imaging biomarkers increasingly connect anatomy and physiology to cognition, development, degeneration, and psychiatric phenotypes.",
        ],
        "limitations": [
            "Scanner, protocol, site, and preprocessing variation can limit reproducibility across cohorts.",
            "Associational imaging biomarkers often need stronger longitudinal or intervention evidence before causal claims are made.",
            "Citation-ranked imaging lists can favor widely reused atlases and datasets over newer mechanistic studies.",
        ],
    },
    {
        "name": "EEG, MEG, and Electrophysiology",
        "slug": "eeg-meg-and-electrophysiology",
        "patterns": [
            "eeg",
            "meg",
            "electroencephal",
            "magnetoencephal",
            "erp",
            "evoked potential",
            "oscillation",
            "spectral",
            "electrophysiology",
            "ecog",
        ],
        "accent": "#0f766e",
        "secondary": "#60a5fa",
        "overview": [
            "Electrophysiology papers organize brain research around fast neural dynamics, oscillations, event-related responses, and disease-related rhythms.",
            "The field is moving toward source localization, multimodal fusion, mobile recordings, and machine-learning assisted decoding.",
            "EEG and MEG remain central where temporal precision, clinical accessibility, or non-invasive monitoring are more important than spatial detail.",
        ],
        "limitations": [
            "Low signal-to-noise ratio, artifacts, and source-localization uncertainty complicate interpretation.",
            "Many studies rely on controlled tasks and may not generalize to everyday behavior or clinical monitoring.",
            "Hardware, montage, preprocessing, and reference choices can make direct comparison difficult.",
        ],
    },
    {
        "name": "Cellular, Molecular, and Synaptic Neuroscience",
        "slug": "cellular-molecular-and-synaptic-neuroscience",
        "patterns": [
            "single cell",
            "single-cell",
            "single unit",
            "neuron",
            "neuronal",
            "synapse",
            "synaptic",
            "glia",
            "microglia",
            "astrocyte",
            "molecular",
            "transcriptomic",
            "calcium imaging",
        ],
        "accent": "#7c3aed",
        "secondary": "#f59e0b",
        "overview": [
            "This area captures cell types, circuits, synaptic plasticity, molecular pathways, and single-cell atlases that explain brain function from the microscopic level.",
            "High-throughput transcriptomics and large-scale cell atlases increasingly bridge molecular identity with anatomy and circuit function.",
            "Classic synaptic and plasticity papers remain prominent because they define mechanisms reused across learning, development, and disease research.",
        ],
        "limitations": [
            "Cellular mechanisms can be difficult to connect directly to whole-brain dynamics or human behavior.",
            "Single-cell sampling, dissociation, alignment, and batch effects can bias inferred cell populations.",
            "Animal and ex vivo findings need careful translation to human disease and cognition.",
        ],
    },
    {
        "name": "Cognitive and Systems Neuroscience",
        "slug": "cognitive-and-systems-neuroscience",
        "patterns": [
            "cognition",
            "memory",
            "attention",
            "perception",
            "decision",
            "learning",
            "behavior",
            "systems neuroscience",
            "hippocampus",
            "prefrontal",
            "visual cortex",
            "auditory cortex",
        ],
        "accent": "#dc2626",
        "secondary": "#f97316",
        "overview": [
            "Cognitive and systems work links brain networks, circuits, and dynamics to perception, action, memory, attention, learning, and decision making.",
            "The literature increasingly combines behavior, neural recording, imaging, computational models, and causal perturbation.",
            "Highly cited papers often become conceptual anchors for how brain systems implement cognition.",
        ],
        "limitations": [
            "Task designs can simplify cognition enough that ecological validity becomes uncertain.",
            "Cross-species alignment between circuits, behavior, and subjective experience remains imperfect.",
            "Correlational neural signatures need causal tests before they are treated as mechanisms.",
        ],
    },
    {
        "name": "Clinical Neurology and Neurodegeneration",
        "slug": "clinical-neurology-and-neurodegeneration",
        "patterns": [
            "alzheimer",
            "parkinson",
            "dementia",
            "stroke",
            "epilepsy",
            "multiple sclerosis",
            "neurology",
            "neurodegeneration",
            "brain tumor",
            "traumatic brain",
            "clinical",
            "patient",
        ],
        "accent": "#be123c",
        "secondary": "#06b6d4",
        "overview": [
            "Clinical brain research concentrates on diagnosis, mechanisms, biomarkers, treatment, and prognosis for neurological and neurodegenerative disorders.",
            "Large cohorts and biomarker frameworks have made dementia, stroke, epilepsy, Parkinson disease, and brain injury especially visible in citation-ranked views.",
            "Translation depends on connecting biological signatures to outcomes that matter to patients and care systems.",
        ],
        "limitations": [
            "Clinical cohorts often differ in disease stage, comorbidity, treatment history, and follow-up duration.",
            "Biomarkers may not transfer cleanly across populations, scanners, care settings, or diagnostic criteria.",
            "High citation counts can favor broad disease frameworks over smaller mechanistic or interventional studies.",
        ],
    },
    {
        "name": "Brain Development, Plasticity, and Connectomics",
        "slug": "brain-development-plasticity-and-connectomics",
        "patterns": [
            "development",
            "developmental",
            "plasticity",
            "connectome",
            "connectivity",
            "network",
            "child",
            "adolescent",
            "aging",
            "lifespan",
            "critical period",
        ],
        "accent": "#0891b2",
        "secondary": "#22c55e",
        "overview": [
            "This category follows how brain structure, networks, and function change across development, learning, aging, and recovery.",
            "Connectomics has shifted the field from isolated regions toward network-level organization and lifespan trajectories.",
            "Plasticity research links cellular mechanisms, experience, rehabilitation, and large-scale brain reorganization.",
        ],
        "limitations": [
            "Developmental and aging studies are sensitive to cohort composition, attrition, and longitudinal sampling intervals.",
            "Network measures can depend heavily on parcellation, thresholding, and acquisition choices.",
            "Plasticity claims need stronger evidence separating transient compensation from durable functional improvement.",
        ],
    },
    {
        "name": "Brain Stimulation, Neurotechnology, and BCI",
        "slug": "brain-stimulation-neurotechnology-and-bci",
        "patterns": [
            "brain stimulation",
            "deep brain stimulation",
            "tms",
            "tdcs",
            "bci",
            "brain-computer",
            "brain machine",
            "neurotechnology",
            "closed-loop",
            "implant",
            "prosthetic",
            "neuromodulation",
        ],
        "accent": "#16a34a",
        "secondary": "#8b5cf6",
        "overview": [
            "Neurotechnology work uses stimulation, implants, decoding, and closed-loop systems to probe or restore brain function.",
            "The field is moving from proof-of-concept control and stimulation toward durable devices, adaptive therapy, and user-centered deployment.",
            "BCI and neuromodulation papers connect engineering performance with clinical function, safety, and usability.",
        ],
        "limitations": [
            "Small cohorts, invasive risk, device maintenance, and long-term stability remain major translation barriers.",
            "Closed-loop effects are hard to separate from placebo, training, medication, and disease fluctuation.",
            "Performance in controlled sessions may not reflect daily-life reliability or user burden.",
        ],
    },
    {
        "name": "Computational Neuroscience and AI",
        "slug": "computational-neuroscience-and-ai",
        "patterns": [
            "computational",
            "model",
            "neural network",
            "machine learning",
            "deep learning",
            "artificial intelligence",
            "bayesian",
            "reinforcement learning",
            "encoding model",
            "decoding",
            "simulation",
        ],
        "accent": "#9333ea",
        "secondary": "#14b8a6",
        "overview": [
            "Computational brain research formalizes neural coding, inference, learning, dynamics, and decoding with statistical and machine-learning models.",
            "AI methods are increasingly used both as analysis tools and as hypotheses for brain computation.",
            "High-impact work often clarifies which computations could plausibly be implemented by neural circuits.",
        ],
        "limitations": [
            "Predictive performance does not necessarily identify causal neural mechanisms.",
            "Models can inherit dataset bias, preprocessing artifacts, and task constraints.",
            "Interpretability and biological plausibility remain central challenges for large AI-based models.",
        ],
    },
    {
        "name": "Cerebrovascular, Metabolism, and Brain Injury",
        "slug": "cerebrovascular-metabolism-and-brain-injury",
        "patterns": [
            "stroke",
            "ischemia",
            "ischaemia",
            "cerebrovascular",
            "blood brain barrier",
            "metabolism",
            "traumatic brain",
            "brain injury",
            "hypoxia",
            "edema",
            "haemorrhage",
            "hemorrhage",
        ],
        "accent": "#ea580c",
        "secondary": "#0284c7",
        "overview": [
            "This category covers vascular, metabolic, inflammatory, and injury pathways that shape brain damage and recovery.",
            "Highly cited studies often define mechanisms, acute care evidence, or biomarkers for stroke, trauma, and barrier dysfunction.",
            "The area is clinically important because small changes in timing, physiology, and treatment windows can alter outcomes.",
        ],
        "limitations": [
            "Acute injury studies can be sensitive to timing, severity, comorbidity, and treatment heterogeneity.",
            "Animal models may not capture human vascular risk, injury complexity, or rehabilitation context.",
            "Translational failures remain common when mechanistic signals are not tied to functional outcomes.",
        ],
    },
    {
        "name": "General Brain Science and Reviews",
        "slug": "general-brain-science-and-reviews",
        "patterns": [
            "review",
            "meta-analysis",
            "systematic review",
            "brain",
            "neuroscience",
            "atlas",
            "framework",
            "consensus",
        ],
        "accent": "#334155",
        "secondary": "#0f766e",
        "overview": [
            "General brain science papers synthesize methods, concepts, atlases, datasets, and cross-domain frameworks.",
            "Reviews and consensus papers are useful entry points because they connect specialized subfields and standardize terminology.",
            "This category also catches broad brain papers that do not fit cleanly into one methodological or disease-focused area.",
        ],
        "limitations": [
            "Broad reviews can dominate citation-ranked lists while obscuring narrower empirical advances.",
            "Taxonomy boundaries are imperfect because many brain papers combine methods, scales, and diseases.",
            "Metadata-level curation cannot replace full-text expert appraisal of claims and evidence quality.",
        ],
    },
]

CATEGORY_BY_NAME = {category["name"]: category for category in CATEGORIES}

KEYWORD_CONVENTION = [
    ("MRI", "Magnetic resonance imaging, structural MRI, diffusion MRI, or MRI-derived brain measures.", "2563eb"),
    ("fMRI", "Functional MRI, BOLD imaging, task activation, resting-state, or functional connectivity.", "7c3aed"),
    ("EEG", "Electroencephalography, event-related potentials, oscillations, or scalp electrophysiology.", "0f766e"),
    ("MEG", "Magnetoencephalography or magnetic source imaging of brain activity.", "0891b2"),
    ("ECoG", "Electrocorticography, intracranial EEG, or cortical surface recordings.", "dc2626"),
    ("single-cell", "Single-cell, single-nucleus, single-unit, spike, transcriptomic, or cell atlas studies.", "f59e0b"),
    ("human", "Human participants, patients, volunteers, human brain tissue, or population cohorts.", "be123c"),
    ("non-human", "Animal, non-human primate, rodent, zebrafish, fly, model organism, or simulation-focused studies.", "a855f7"),
    ("connectome", "Brain connectivity, connectomics, networks, tractography, or atlas-based mapping.", "16a34a"),
    ("stimulation", "Brain stimulation, DBS, TMS, tDCS, neuromodulation, implants, or closed-loop intervention.", "ea580c"),
]
KEYWORD_COLORS = {keyword: color for keyword, _, color in KEYWORD_CONVENTION}

RECOGNIZED_VENUES = [
    "Nature",
    "Science",
    "Cell",
    "Neuron",
    "The Lancet",
    "Nature Neuroscience",
    "Nature Reviews Neuroscience",
    "Nature Reviews Neurology",
    "PNAS",
    "Proceedings of the National Academy of Sciences",
    "Brain",
    "Cerebral Cortex",
    "NeuroImage",
    "Journal of Neuroscience",
    "Annals of Neurology",
    "Neurology",
    "eLife",
]

CSV_FIELDS = [
    "rank",
    "year",
    "title",
    "authors",
    "venue",
    "publicationDate",
    "citationCount",
    "influentialCitationCount",
    "importanceScore",
    "category",
    "keywordTags",
    "keyIdea",
    "strengths",
    "limitations",
    "doi",
    "sourceId",
    "url",
    "semanticScholarUrl",
    "openAccessPdf",
    "githubUrl",
    "workType",
    "language",
    "relevanceReason",
]

CANDIDATE_FIELDS = [
    "year",
    "title",
    "authors",
    "venue",
    "publicationDate",
    "citationCount",
    "category",
    "keywordTags",
    "doi",
    "sourceId",
    "url",
    "workType",
    "language",
    "relevanceScore",
    "relevanceReason",
    "abstractSnippet",
]


def ensure_dirs():
    for path in [DATA_DIR, DOCS_DIR, PAPER_DIR, CACHE_DIR, DOCS_DIR / "data", DOCS_DIR / "assets", DOCS_DIR / "assets" / "taxonomy", DOCS_DIR / "paper"]:
        path.mkdir(parents=True, exist_ok=True)


def clean_text(value):
    if value is None:
        return ""
    text = str(value)
    text = text.replace("\u2010", "-").replace("\u2011", "-").replace("\u2012", "-").replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
    return re.sub(r"\s+", " ", text).strip()


def title_key(value):
    text = clean_text(value).lower()
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def clean_github_url(value):
    text = clean_text(value)
    match = re.search(r"https?://github\.com/[^\s<>\"')]+", text, flags=re.I)
    if not match:
        return ""
    path = match.group(0).split("github.com/", 1)[1].rstrip(".,;")
    return f"https://github.com/{path}"


def semantic_scholar_url(row):
    paper_id = clean_text(row.get("paperId") or row.get("sourceId"))
    if re.fullmatch(r"[0-9a-f]{40}", paper_id, flags=re.I):
        return f"https://www.semanticscholar.org/paper/{paper_id}"
    return clean_text(row.get("semanticScholarUrl"))


def is_official_github_entry(value):
    return isinstance(value, dict) and bool(value.get("githubOfficial")) and bool(value.get("mentionedInPaper"))


def official_github_url_from_paper_text(row):
    official_context = re.compile(
        r"\b(code|codes|source|implementation|repo|repository|software|toolbox|package|model|models|weights|data|dataset|project|available|released|open-source|publicly)\b",
        flags=re.I,
    )
    for field in ("abstract", "abstractSnippet"):
        text = clean_text(row.get(field))
        for match in re.finditer(r"https?://github\.com/[^\s<>\"')]+", text, flags=re.I):
            window = text[max(0, match.start() - 160): match.end() + 80]
            if official_context.search(window):
                return clean_github_url(match.group(0))
    return ""


def load_github_links():
    path = DATA_DIR / GITHUB_LINKS_JSON
    if not path.exists():
        return {}
    payload = json.loads(path.read_text(encoding="utf-8"))
    links = {}
    for key, value in payload.get("links", {}).items():
        if not is_official_github_entry(value):
            continue
        github_url = clean_github_url(value.get("githubUrl") if isinstance(value, dict) else value)
        if github_url:
            links[key] = {**value, "githubUrl": github_url} if isinstance(value, dict) else {"githubUrl": github_url}
    return links


def apply_github_links(rows):
    links = load_github_links()
    for row in rows:
        match = links.get(title_key(row.get("title")))
        if match and is_official_github_entry(match):
            row["githubUrl"] = clean_github_url(match.get("githubUrl", ""))
        else:
            row["githubUrl"] = ""
    return rows


def strip_doi(doi):
    doi = clean_text(doi)
    doi = re.sub(r"^https?://(dx\.)?doi\.org/", "", doi, flags=re.I)
    return doi


def slugify(text):
    text = clean_text(text).lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    return text.strip("-") or "item"


def abstract_from_inverted_index(index):
    if not index:
        return ""
    positions = []
    for word, offsets in index.items():
        for offset in offsets:
            positions.append((offset, word))
    positions.sort()
    return clean_text(" ".join(word for _, word in positions))


def first_sentence(text):
    text = clean_text(text)
    if not text:
        return ""
    match = re.search(r"(.{40,420}?[.!?])\s", text + " ")
    if match:
        return match.group(1).strip()
    return text[:360].strip()


def authors_from_work(work, limit=6):
    authors = []
    for authorship in work.get("authorships") or []:
        author = authorship.get("author") or {}
        name = clean_text(author.get("display_name") or authorship.get("raw_author_name"))
        if name:
            authors.append(name)
    if not authors:
        return "Unknown authors"
    if len(authors) > limit:
        return ", ".join(authors[:limit]) + ", et al."
    return ", ".join(authors)


def venue_from_work(work):
    primary = work.get("primary_location") or {}
    source = primary.get("source") or {}
    venue = clean_text(source.get("display_name") or primary.get("raw_source_name"))
    return venue or "Unknown venue"


def link_from_work(work):
    doi = strip_doi(work.get("doi"))
    if doi:
        return f"https://doi.org/{doi}"
    primary = work.get("primary_location") or {}
    if primary.get("landing_page_url"):
        return clean_text(primary["landing_page_url"])
    ids = work.get("ids") or {}
    for key in ["pmid", "pmcid"]:
        if ids.get(key):
            return clean_text(ids[key])
    return clean_text(work.get("id"))


def pdf_from_work(work):
    best = work.get("best_oa_location") or {}
    primary = work.get("primary_location") or {}
    open_access = work.get("open_access") or {}
    return clean_text(best.get("pdf_url") or primary.get("pdf_url") or open_access.get("oa_url"))


def concept_terms(work):
    terms = []
    for concept in work.get("concepts") or []:
        name = clean_text(concept.get("display_name"))
        score = concept.get("score") or 0
        if name:
            terms.append((name, score))
    for keyword in work.get("keywords") or []:
        name = clean_text(keyword.get("display_name"))
        score = keyword.get("score") or 0
        if name:
            terms.append((name, score))
    return terms


def relevant_text(work, abstract=""):
    title = clean_text(work.get("title") or work.get("display_name"))
    concepts = " ".join(name for name, score in concept_terms(work) if score >= 0.12)
    return f"{title} {abstract} {concepts}".lower()


def relevance_score(work, abstract):
    text = relevant_text(work, abstract)
    score = 0
    reasons = []
    for term in RELEVANCE_TERMS:
        if term in text:
            score += 1
            if len(reasons) < 8:
                reasons.append(term)
    for concept, concept_score in concept_terms(work):
        if concept.lower() == "neuroscience" and concept_score >= 0.2:
            score += 3
            reasons.append("neuroscience concept")
        elif concept_score >= 0.35 and any(term in concept.lower() for term in ["brain", "neuro", "cortex", "neuron", "synapse", "connect"]):
            score += 2
            reasons.append(concept)
    if "brain" in clean_text(work.get("title") or "").lower():
        score += 3
        reasons.append("brain in title")
    return score, "; ".join(dict.fromkeys(reasons)) or "brain and neuroscience query"


def assign_category(text):
    text = text.lower()
    best_name = "General Brain Science and Reviews"
    best_score = 0
    for category in CATEGORIES:
        score = sum(1 for pattern in category["patterns"] if pattern in text)
        if category["name"] == "General Brain Science and Reviews":
            score -= 1
        if score > best_score:
            best_name = category["name"]
            best_score = score
    return best_name


def assign_keywords(text):
    lower = text.lower()
    tags = []
    checks = {
        "MRI": [" mri", "magnetic resonance", "diffusion tensor", "diffusion mri", "structural mri"],
        "fMRI": ["fmri", "functional magnetic resonance", "bold", "resting-state", "functional connectivity"],
        "EEG": ["eeg", "electroencephal", "event-related potential", "erp", "scalp potential"],
        "MEG": ["meg", "magnetoencephal"],
        "ECoG": ["ecog", "electrocortic", "intracranial eeg", "subdural"],
        "single-cell": ["single-cell", "single cell", "single nucleus", "single-nucleus", "single unit", "spike sorting", "transcriptomic", "cell atlas"],
        "human": ["human", "patient", "volunteer", "participant", "cohort", "children", "adolescent", "adult", "clinical"],
        "non-human": ["mouse", "mice", "rat", "rodent", "monkey", "macaque", "primate", "zebrafish", "drosophila", "animal", "murine", "non-human"],
        "connectome": ["connectome", "connectivity", "network", "tractography", "atlas", "parcellation"],
        "stimulation": ["stimulation", "deep brain stimulation", "dbs", "tms", "tdcs", "neuromodulation", "closed-loop", "implant"],
    }
    padded = " " + lower
    for keyword, needles in checks.items():
        if any(needle in padded for needle in needles):
            tags.append(keyword)
    return tags


def recognized_venue(venue):
    venue_lower = venue.lower()
    return any(candidate.lower() in venue_lower for candidate in RECOGNIZED_VENUES)


def work_type_label(work_type):
    return clean_text(work_type).replace("-", " ") or "work"


def importance_score(row):
    score = math.log1p(row["citationCount"]) * 18
    score += min(row.get("relevanceScore", 0), 20) * 2
    if recognized_venue(row["venue"]):
        score += 12
    if row["workType"] in {"review", "meta-analysis"} or "review" in row["title"].lower():
        score += 8
    if row.get("openAccessPdf"):
        score += 3
    if row["keywordTags"]:
        score += len(row["keywordTags"].split(";")) * 1.2
    return round(score, 2)


def strengths_for(row):
    strengths = [f"high citation signal ({row['citationCount']:,})"]
    tags = {tag for tag in row["keywordTags"].split(";") if tag}
    if recognized_venue(row["venue"]):
        strengths.append("recognized venue")
    if row.get("openAccessPdf"):
        strengths.append("open-access PDF metadata")
    if "review" in row["workType"] or "review" in row["title"].lower():
        strengths.append("synthesis or review value")
    if "human" in tags:
        strengths.append("human evidence signal")
    if "single-cell" in tags:
        strengths.append("cell-resolution signal")
    if not strengths:
        strengths.append("selected by citation count from the audited brain candidate pool")
    return "; ".join(dict.fromkeys(strengths[:4]))


def limitations_for(row):
    category = CATEGORY_BY_NAME[row["category"]]
    limitations = list(category["limitations"][:2])
    tags = set(row["keywordTags"].split(";")) if row["keywordTags"] else set()
    if "human" in tags and "non-human" not in tags:
        limitations.append("Human cohort effects can depend on recruitment, demographics, comorbidity, and measurement context.")
    elif "non-human" in tags:
        limitations.append("Non-human findings need careful translation to human brain organization and clinical outcomes.")
    elif "MRI" in tags or "fMRI" in tags:
        limitations.append("Imaging conclusions can depend on acquisition, preprocessing, and model specification.")
    else:
        limitations.append("Metadata-level ranking should be complemented with full-text expert review before strong claims are made.")
    return "; ".join(dict.fromkeys(limitations[:3]))


def normalize_work(work):
    abstract = abstract_from_inverted_index(work.get("abstract_inverted_index"))
    score, reason = relevance_score(work, abstract)
    title = clean_text(work.get("title") or work.get("display_name"))
    venue = venue_from_work(work)
    concept_text_value = " ".join(name for name, _ in concept_terms(work))
    aggregate_text = f"{title} {abstract} {venue} {concept_text_value}"
    category = assign_category(aggregate_text)
    tags = assign_keywords(aggregate_text)
    doi = strip_doi(work.get("doi"))
    row = {
        "rank": "",
        "year": int(work.get("publication_year") or 0),
        "title": title,
        "authors": authors_from_work(work),
        "venue": venue,
        "publicationDate": clean_text(work.get("publication_date")) or str(work.get("publication_year") or ""),
        "citationCount": int(work.get("cited_by_count") or 0),
        "influentialCitationCount": 0,
        "category": category,
        "keywordTags": ";".join(tags),
        "doi": doi,
        "sourceId": clean_text(work.get("id")),
        "url": link_from_work(work),
        "openAccessPdf": pdf_from_work(work),
        "workType": work_type_label(work.get("type")),
        "language": clean_text(work.get("language")) or "unknown",
        "relevanceScore": score,
        "relevanceReason": reason,
        "abstract": abstract,
        "abstractSnippet": first_sentence(abstract) or "",
    }
    row["keyIdea"] = first_sentence(abstract) or f"Positions {title} within {category}."
    row["strengths"] = strengths_for(row)
    row["limitations"] = limitations_for(row)
    row["importanceScore"] = importance_score(row)
    return row


def cache_path_for_year(year):
    return CACHE_DIR / f"s2_brain_candidates_{year}.json"


def s2_paper_key(paper):
    ext = paper.get("externalIds") or {}
    for key in ["DOI", "ArXiv", "PubMed", "CorpusId"]:
        if ext.get(key):
            return f"{key}:{ext[key]}".lower()
    return (paper.get("paperId") or paper.get("url") or paper.get("title") or "").lower()


def s2_authors(paper, limit=6):
    names = [clean_text(author.get("name")) for author in paper.get("authors") or [] if author.get("name")]
    if not names:
        return "Unknown authors"
    if len(names) > limit:
        return ", ".join(names[:limit]) + ", et al."
    return ", ".join(names)


def s2_venue(paper):
    venue = clean_text(paper.get("venue"))
    publication_venue = paper.get("publicationVenue") or {}
    return venue or clean_text(publication_venue.get("name")) or "Unknown venue"


def s2_fields_text(paper):
    fields = paper.get("s2FieldsOfStudy") or []
    categories = [clean_text(field.get("category")) for field in fields if field.get("category")]
    types = [clean_text(item) for item in paper.get("publicationTypes") or []]
    return " ".join(categories + types)


def s2_relevance_score(paper):
    title = clean_text(paper.get("title"))
    abstract = clean_text(paper.get("abstract"))
    text = f"{title} {abstract} {s2_fields_text(paper)}".lower()
    if ("brain natriuretic peptide" in text or "nt-probnp" in text or "b-type natriuretic" in text) and not any(
        term in text for term in ["neuron", "cortex", "cerebral", "neuro", "mri", "fmri", "eeg", "stroke", "dementia"]
    ):
        return 0, "excluded natriuretic-peptide usage"
    score = 0
    reasons = []
    for term in RELEVANCE_TERMS:
        if term in text:
            score += 1
            if len(reasons) < 8:
                reasons.append(term)
    if "brain" in title.lower():
        score += 3
        reasons.append("brain in title")
    if any(term in text for term in ["neuroscience", "neurology", "neuroimage", "neuron", "cortex", "hippocampus", "synapse"]):
        score += 3
        reasons.append("neuroscience term")
    if not reasons and "Medicine" in s2_fields_text(paper):
        return 0, "weak brain relevance"
    return score, "; ".join(dict.fromkeys(reasons)) or "Semantic Scholar brain query"


def normalize_s2_paper(paper):
    abstract = clean_text(paper.get("abstract"))
    score, reason = s2_relevance_score(paper)
    title = clean_text(paper.get("title"))
    venue = s2_venue(paper)
    aggregate_text = f"{title} {abstract} {venue} {s2_fields_text(paper)}"
    category = assign_category(aggregate_text)
    tags = assign_keywords(aggregate_text)
    ext = paper.get("externalIds") or {}
    oa = paper.get("openAccessPdf") or {}
    doi = strip_doi(ext.get("DOI", ""))
    url = f"https://doi.org/{doi}" if doi else clean_text(paper.get("url"))
    row = {
        "rank": "",
        "year": int(paper.get("year") or 0),
        "title": title,
        "authors": s2_authors(paper),
        "venue": venue,
        "publicationDate": clean_text(paper.get("publicationDate")) or str(paper.get("year") or ""),
        "citationCount": int(paper.get("citationCount") or 0),
        "influentialCitationCount": int(paper.get("influentialCitationCount") or 0),
        "category": category,
        "keywordTags": ";".join(tags),
        "doi": doi,
        "sourceId": clean_text(paper.get("paperId")),
        "url": url,
        "semanticScholarUrl": semantic_scholar_url({"sourceId": paper.get("paperId", "")}),
        "openAccessPdf": clean_text(oa.get("url")) if isinstance(oa, dict) else "",
        "workType": "; ".join(clean_text(item) for item in paper.get("publicationTypes") or []) or "paper",
        "language": "unknown",
        "relevanceScore": score,
        "relevanceReason": reason,
        "abstract": abstract,
        "abstractSnippet": first_sentence(abstract) or "",
    }
    row["keyIdea"] = first_sentence(abstract) or f"Positions {title} within {category}."
    row["strengths"] = strengths_for(row)
    row["limitations"] = limitations_for(row)
    row["importanceScore"] = importance_score(row)
    return row


def refresh_row_enrichment(row):
    row = dict(row)
    row["year"] = int(row.get("year") or 0)
    row["citationCount"] = int(row.get("citationCount") or 0)
    row["influentialCitationCount"] = int(row.get("influentialCitationCount") or 0)
    row["importanceScore"] = importance_score(row)
    row["strengths"] = strengths_for(row)
    row["limitations"] = limitations_for(row)
    if not row.get("keyIdea"):
        row["keyIdea"] = f"Positions {row.get('title', 'this paper')} within {row.get('category', 'brain research')}."
    return row


def fetch_s2_year_query(session, year, query, max_pages=S2_MAX_PAGES_PER_QUERY):
    params = {
        "query": query,
        "year": str(year),
        "fields": S2_FIELDS,
        "sort": "citationCount:desc",
    }
    papers = []
    token = None
    for _ in range(max_pages):
        if token:
            params["token"] = token
        for attempt in range(1, 6):
            response = session.get(S2_BULK_URL, params=params, timeout=90)
            if response.status_code == 429:
                wait = 12 * attempt
                print(f"[s2] 429 for {year} {query}; retrying in {wait}s", flush=True)
                time.sleep(wait)
                continue
            if response.status_code in {500, 502, 503, 504}:
                wait = min(45, 2 ** attempt)
                print(f"[s2] {response.status_code} for {year} {query}; retrying in {wait}s", flush=True)
                time.sleep(wait)
                continue
            response.raise_for_status()
            data = response.json()
            papers.extend(data.get("data") or [])
            token = data.get("token")
            break
        else:
            break
        if not token:
            break
        time.sleep(1.0)
    return papers


def collect_year(session, year, refresh=False):
    cache_path = cache_path_for_year(year)
    if cache_path.exists() and not refresh:
        rows = [refresh_row_enrichment(row) for row in json.loads(cache_path.read_text(encoding="utf-8"))]
        cache_path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
        return rows
    merged = {}
    for query in S2_QUERIES:
        print(f"[collect] {year} :: {query}", flush=True)
        try:
            for paper in fetch_s2_year_query(session, year, query):
                if paper.get("year") != year or not paper.get("title"):
                    continue
                row = normalize_s2_paper(paper)
                if row["relevanceScore"] <= 0:
                    continue
                merged[s2_paper_key(paper)] = row
        except Exception as exc:
            print(f"[warn] {year} {query}: {exc}", flush=True)
        if len(merged) >= CANDIDATES_PER_YEAR:
            break
        time.sleep(1.0)
    rows = sorted(merged.values(), key=lambda item: (-item["citationCount"], -item["relevanceScore"], item["title"].lower()))
    rows = rows[:CANDIDATES_PER_YEAR]
    cache_path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
    return rows


def collect_papers(refresh=False):
    session = requests.Session()
    session.headers.update({"User-Agent": "awesome-brain-curation/1.0"})
    all_candidates = []
    selected = []
    for year in YEARS:
        rows = collect_year(session, year, refresh=refresh)
        for row in rows:
            all_candidates.append(row)
        year_selected = sorted(rows, key=lambda item: (-item["citationCount"], -item["relevanceScore"], item["title"].lower()))[:TARGET_PER_YEAR]
        for rank, row in enumerate(year_selected, start=1):
            row = dict(row)
            row["rank"] = rank
            selected.append(row)
        print(f"[year {year}] candidates={len(rows)} selected={len(year_selected)}", flush=True)
    selected.sort(key=lambda item: (item["year"], item["rank"]))
    all_candidates.sort(key=lambda item: (item["year"], -item["citationCount"], item["title"].lower()))
    return selected, all_candidates


def selected_csv_row(row):
    return {field: row.get(field, "") for field in CSV_FIELDS}


def candidate_csv_row(row):
    return {field: row.get(field, "") for field in CANDIDATE_FIELDS}


def write_csv(path, rows, fields):
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def write_json_csv(selected, candidates):
    DATA_DIR.mkdir(exist_ok=True)
    apply_github_links(selected)
    selected_export = [selected_csv_row(row) | {"abstract": row.get("abstract", "")} for row in selected]
    candidate_export = [candidate_csv_row(row) for row in candidates]
    (DATA_DIR / PAPERS_JSON).write_text(json.dumps(selected_export, ensure_ascii=False, indent=2), encoding="utf-8")
    (DATA_DIR / CANDIDATES_JSON).write_text(json.dumps(candidate_export, ensure_ascii=False, indent=2), encoding="utf-8")
    write_csv(DATA_DIR / PAPERS_CSV, selected_export, CSV_FIELDS + ["abstract"])
    write_csv(DATA_DIR / CANDIDATES_CSV, candidate_export, CANDIDATE_FIELDS)
    for year in YEARS:
        year_selected = [selected_csv_row(row) | {"abstract": row.get("abstract", "")} for row in selected if row["year"] == year]
        year_candidates = [candidate_csv_row(row) for row in candidates if row["year"] == year]
        write_csv(DATA_DIR / f"papers_{year}.csv", year_selected, CSV_FIELDS + ["abstract"])
        write_csv(DATA_DIR / f"candidates_top{CANDIDATES_PER_YEAR}_{year}.csv", year_candidates, CANDIDATE_FIELDS)


def write_taxonomy_dataset(selected):
    fields = [
        "rank",
        "year",
        "category",
        "keywordTags",
        "title",
        "authors",
        "venue",
        "citationCount",
        "importanceScore",
        "keyIdea",
        "strengths",
        "limitations",
        "url",
        "doi",
        "sourceId",
    ]
    write_csv(DATA_DIR / TAXONOMY_CSV, selected, fields)


def year_stats(rows):
    stats = {}
    for year in YEARS:
        subset = [row for row in rows if row["year"] == year]
        if not subset:
            continue
        top = max(subset, key=lambda item: item["citationCount"])
        stats[year] = {
            "count": len(subset),
            "citations": sum(row["citationCount"] for row in subset),
            "top": {"title": top["title"], "url": top["url"], "citations": top["citationCount"]},
        }
    return stats


def category_stats(rows):
    counts = Counter(row["category"] for row in rows)
    citations = defaultdict(int)
    for row in rows:
        citations[row["category"]] += row["citationCount"]
    return counts, citations


def keyword_stats(rows):
    counts = Counter()
    for row in rows:
        for tag in row["keywordTags"].split(";"):
            if tag:
                counts[tag] += 1
    return counts


def research_overview_html():
    return """
    <section class="research-brief" id="researchBrief" aria-labelledby="research-timeline-title">
      <h2 id="research-timeline-title">Research Timeline</h2>
      <div class="timeline-copy">
        <p>1900-2026년 brain 연구 코퍼스는 해부학, 병리학, 생리학 중심의 초기 뇌과학에서 출발해 임상 신경학, 세포·분자 신경과학, 신경영상, 전기생리, connectomics, 계산신경과학과 neurotechnology로 확장된 장기 지형도다. 인용 기반으로 선별된 12,700편은 특정 연도별 유행보다, 후속 연구의 공통 언어가 된 atlas, 측정법, 임상 기준, 리뷰, 질환 코호트, 분석 도구의 축적을 강하게 드러낸다.</p>
        <p>가장 큰 흐름은 일반 뇌과학과 리뷰, 임상 신경학·신경퇴행, 세포·분자·시냅스 연구가 넓은 기반을 형성하고, 그 위에 인지·시스템 신경과학, 발달·가소성·connectomics, EEG/MEG, 신경영상, 뇌혈관·손상, 계산신경과학과 BCI가 접속하는 구조다. 최근 구간으로 올수록 단일 방법론보다 다중모달 데이터, 장기 추적, 네트워크 수준 해석, AI 기반 분석, 자극·폐루프 시스템을 결합하는 연구가 더 중요해진다.</p>
      </div>
      <h2>Research Insights</h2>
      <div class="research-insights">
        <article class="insight-box">
          <div class="insight-label">Infrastructure</div>
          <h3>공유 인프라가 연구 방향을 재편한다</h3>
          <p>높은 인용 신호는 단일 발견뿐 아니라 atlas, 좌표계, 영상·전기생리 프로토콜, 임상 척도처럼 여러 하위 분야가 함께 쓰는 기반 논문에 집중된다.</p>
          <p class="insight-implication">시사점: 새 연구는 독립 결과보다 재사용 가능한 데이터, 표준, 도구를 동반할 때 장기 영향력이 커진다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Clinical Translation</div>
          <h3>질병 연구는 회로·분자·영상의 결합으로 이동한다</h3>
          <p>신경퇴행, 뇌손상, 종양, 뇌혈관 연구는 병리 분류에서 바이오마커, 영상, 세포 기전, 치료 반응을 함께 읽는 방향으로 넓어지고 있다.</p>
          <p class="insight-implication">시사점: 임상적 유용성은 기전 설명과 환자군 재현성을 동시에 요구한다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Multi-scale Models</div>
          <h3>발달·가소성·connectomics가 스케일을 연결한다</h3>
          <p>세포와 시냅스, 국소 회로, 전뇌 네트워크, 행동 변화를 이어 보려는 연구가 증가하며 생애주기와 질환 진행을 함께 설명하는 축이 된다.</p>
          <p class="insight-implication">시사점: 횡단면 결과만으로는 부족하며 장기 추적과 네트워크 검증이 핵심 병목이다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">AI And Computation</div>
          <h3>계산 모델은 관찰을 예측 문제로 바꾼다</h3>
          <p>계산신경과학과 AI는 뇌 신호를 해석하는 보조 도구에서 representation, decoding, disease trajectory 예측을 다루는 연구 언어로 확장된다.</p>
          <p class="insight-implication">시사점: 모델 성능보다 해석 가능성, 외부 검증, 데이터 편향 통제가 중요해진다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Neurotechnology</div>
          <h3>자극·BCI는 임상 전환의 관문이 된다</h3>
          <p>뇌자극, neurotechnology, BCI 연구는 신호 측정에서 개입과 폐루프 제어로 이동하며 안전성, 장기 안정성, 개인화가 핵심 기준이 된다.</p>
          <p class="insight-implication">시사점: 실제 적용을 위해서는 정확도와 함께 사용성, 위험 관리, 환자별 적응 전략이 필요하다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Open Gaps</div>
          <h3>표준화·인과성·다양성이 남은 병목이다</h3>
          <p>인용 기반 지도는 강한 축을 보여주지만, 인구집단 다양성, 장비·분석 표준화, 인과 검증, 최신 연구의 인용 지연은 여전히 해석상의 주의점이다.</p>
          <p class="insight-implication">시사점: 다음 단계의 brain research map은 citation rank와 함께 reproducibility와 clinical utility를 함께 평가해야 한다.</p>
        </article>
      </div>
    </section>
"""


def research_copy():
    return {
        "en": """
      <h2 id="research-timeline-title">Research Timeline</h2>
      <div class="timeline-copy">
        <p>The 1900-2026 brain corpus maps a long shift from anatomy, pathology, and physiology toward clinical neurology, cellular and molecular neuroscience, neuroimaging, electrophysiology, connectomics, computational neuroscience, and neurotechnology. The 12,700 citation-ranked papers emphasize shared research infrastructure: atlases, measurement protocols, clinical criteria, reviews, disease cohorts, and analysis tools that became common language across the field.</p>
        <p>The largest bodies of work are general brain science and reviews, clinical neurology and neurodegeneration, and cellular, molecular, and synaptic neuroscience. Around them, cognitive and systems neuroscience, development and plasticity, EEG/MEG, neuroimaging, vascular and injury research, computational neuroscience, and BCI form a multi-scale map. Recent work increasingly combines multimodal data, longitudinal cohorts, network-level interpretation, AI-assisted analysis, and closed-loop intervention.</p>
      </div>
      <h2>Research Insights</h2>
      <div class="research-insights">
        <article class="insight-box"><div class="insight-label">Infrastructure</div><h3>Shared infrastructure reshapes the field</h3><p>Highly cited work often defines reusable atlases, coordinates, imaging and electrophysiology protocols, or clinical scales rather than a single isolated finding.</p><p class="insight-implication">Implication: durable impact grows when new studies ship reusable data, standards, and tools.</p></article>
        <article class="insight-box"><div class="insight-label">Clinical Translation</div><h3>Disease research is becoming multi-modal</h3><p>Neurodegeneration, brain injury, tumor, and vascular studies increasingly combine pathology, biomarkers, imaging, cell mechanisms, and treatment response.</p><p class="insight-implication">Implication: clinical usefulness depends on both mechanistic explanation and reproducible patient cohorts.</p></article>
        <article class="insight-box"><div class="insight-label">Multi-scale Models</div><h3>Plasticity and connectomics link scales</h3><p>Development and plasticity research connects cells, synapses, local circuits, whole-brain networks, and behavior across lifespan and disease progression.</p><p class="insight-implication">Implication: longitudinal evidence and network validation are central bottlenecks.</p></article>
        <article class="insight-box"><div class="insight-label">AI And Computation</div><h3>Computation turns observation into prediction</h3><p>Computational neuroscience and AI expand from analysis aids into languages for representation, decoding, and disease-trajectory prediction.</p><p class="insight-implication">Implication: interpretability, external validation, and bias control matter as much as model accuracy.</p></article>
        <article class="insight-box"><div class="insight-label">Neurotechnology</div><h3>Stimulation and BCI mark the translation frontier</h3><p>Brain stimulation, neurotechnology, and BCI move from measurement toward intervention and closed-loop control.</p><p class="insight-implication">Implication: deployment needs usability, safety management, and patient-specific adaptation.</p></article>
        <article class="insight-box"><div class="insight-label">Open Gaps</div><h3>Standardization, causality, and diversity remain open</h3><p>Citation maps reveal strong axes, but population diversity, protocol standardization, causal validation, and citation lag still shape interpretation.</p><p class="insight-implication">Implication: future maps should combine citation rank with reproducibility and clinical utility.</p></article>
      </div>
""",
        "ko": """
      <h2 id="research-timeline-title">연구 타임라인</h2>
      <div class="timeline-copy">
        <p>1900-2026년 brain 연구 코퍼스는 해부학, 병리학, 생리학 중심의 초기 뇌과학에서 출발해 임상 신경학, 세포·분자 신경과학, 신경영상, 전기생리, connectomics, 계산신경과학과 neurotechnology로 확장된 장기 지형도다. 인용 기반으로 선별된 12,700편은 특정 연도별 유행보다, 후속 연구의 공통 언어가 된 atlas, 측정법, 임상 기준, 리뷰, 질환 코호트, 분석 도구의 축적을 강하게 드러낸다.</p>
        <p>가장 큰 흐름은 일반 뇌과학과 리뷰, 임상 신경학·신경퇴행, 세포·분자·시냅스 연구가 넓은 기반을 형성하고, 그 위에 인지·시스템 신경과학, 발달·가소성·connectomics, EEG/MEG, 신경영상, 뇌혈관·손상, 계산신경과학과 BCI가 접속하는 구조다. 최근 구간으로 올수록 단일 방법론보다 다중모달 데이터, 장기 추적, 네트워크 수준 해석, AI 기반 분석, 자극·폐루프 시스템을 결합하는 연구가 더 중요해진다.</p>
      </div>
      <h2>연구 인사이트</h2>
      <div class="research-insights">
        <article class="insight-box"><div class="insight-label">Infrastructure</div><h3>공유 인프라가 연구 방향을 재편한다</h3><p>높은 인용 신호는 단일 발견뿐 아니라 atlas, 좌표계, 영상·전기생리 프로토콜, 임상 척도처럼 여러 하위 분야가 함께 쓰는 기반 논문에 집중된다.</p><p class="insight-implication">시사점: 새 연구는 독립 결과보다 재사용 가능한 데이터, 표준, 도구를 동반할 때 장기 영향력이 커진다.</p></article>
        <article class="insight-box"><div class="insight-label">Clinical Translation</div><h3>질병 연구는 회로·분자·영상의 결합으로 이동한다</h3><p>신경퇴행, 뇌손상, 종양, 뇌혈관 연구는 병리 분류에서 바이오마커, 영상, 세포 기전, 치료 반응을 함께 읽는 방향으로 넓어지고 있다.</p><p class="insight-implication">시사점: 임상적 유용성은 기전 설명과 환자군 재현성을 동시에 요구한다.</p></article>
        <article class="insight-box"><div class="insight-label">Multi-scale Models</div><h3>발달·가소성·connectomics가 스케일을 연결한다</h3><p>세포와 시냅스, 국소 회로, 전뇌 네트워크, 행동 변화를 이어 보려는 연구가 증가하며 생애주기와 질환 진행을 함께 설명하는 축이 된다.</p><p class="insight-implication">시사점: 횡단면 결과만으로는 부족하며 장기 추적과 네트워크 검증이 핵심 병목이다.</p></article>
        <article class="insight-box"><div class="insight-label">AI And Computation</div><h3>계산 모델은 관찰을 예측 문제로 바꾼다</h3><p>계산신경과학과 AI는 뇌 신호를 해석하는 보조 도구에서 representation, decoding, disease trajectory 예측을 다루는 연구 언어로 확장된다.</p><p class="insight-implication">시사점: 모델 성능보다 해석 가능성, 외부 검증, 데이터 편향 통제가 중요해진다.</p></article>
        <article class="insight-box"><div class="insight-label">Neurotechnology</div><h3>자극·BCI는 임상 전환의 관문이 된다</h3><p>뇌자극, neurotechnology, BCI 연구는 신호 측정에서 개입과 폐루프 제어로 이동하며 안전성, 장기 안정성, 개인화가 핵심 기준이 된다.</p><p class="insight-implication">시사점: 실제 적용을 위해서는 정확도와 함께 사용성, 위험 관리, 환자별 적응 전략이 필요하다.</p></article>
        <article class="insight-box"><div class="insight-label">Open Gaps</div><h3>표준화·인과성·다양성이 남은 병목이다</h3><p>인용 기반 지도는 강한 축을 보여주지만, 인구집단 다양성, 장비·분석 표준화, 인과 검증, 최신 연구의 인용 지연은 여전히 해석상의 주의점이다.</p><p class="insight-implication">시사점: 다음 단계의 brain research map은 citation rank와 함께 reproducibility와 clinical utility를 함께 평가해야 한다.</p></article>
      </div>
""",
    }


def overall_research_templates():
    return {
        "en": {
            "timelineTitle": "Research Timeline",
            "summary": [
                "For {range}, this brain corpus contains {papers} selected papers across {activeYears} active years, with {citations} citations. The strongest taxonomy signals are {topCategories}, and the most active year is {peakYear} ({peakYearCount} papers).",
                "The leading citation-ranked paper is \"{topPaper}\" ({topPaperYear}, {topPaperCitations} citations) in {topPaperCategory}. Keywords such as {topKeywords} show how the period connects measurement, mechanisms, clinical translation, and computation.",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {
                    "label": "Period Shape",
                    "title": "The selected range changes the field map",
                    "body": "{topCategory} accounts for {topCategoryCount} papers, so the visible corpus is anchored by the taxonomies that were strongest in {range}.",
                    "implication": "Implication: compare adjacent periods before treating one taxonomy as the field's long-term center.",
                },
                {
                    "label": "Citation Mass",
                    "title": "Citation concentration identifies shared infrastructure",
                    "body": "The range carries {citations} citations, with the citation peak around {peakCitationYear}. Highly cited papers often define reusable atlases, protocols, cohorts, or analysis tools.",
                    "implication": "Implication: durable brain research impact often comes from resources that other subfields can reuse.",
                },
                {
                    "label": "Translation",
                    "title": "Clinical and mechanistic signals should be read together",
                    "body": "The top paper, \"{topPaper}\", sits in {topPaperCategory}, while leading categories include {topCategories}. This mix shows whether the period is driven by disease cohorts, mechanisms, measurement, or computation.",
                    "implication": "Implication: strong summaries should connect patient relevance with causal or measurement evidence.",
                },
                {
                    "label": "Methods",
                    "title": "Keywords expose the period's methodological spine",
                    "body": "Frequent tags such as {topKeywords} indicate which instruments, models, or research settings organize the selected years.",
                    "implication": "Implication: keyword shifts are useful early signals before citation counts fully mature.",
                },
                {
                    "label": "Open Gaps",
                    "title": "Citation-ranked maps still need expert interpretation",
                    "body": "Recent years, underrepresented populations, null results, and protocol differences may be muted even when the period summary looks stable.",
                    "implication": "Implication: use this view as a navigation layer, then pair it with full-text review and reproducibility checks.",
                },
            ],
        },
        "ko": {
            "timelineTitle": "연구 타임라인",
            "summary": [
                "{range} 기간의 brain 코퍼스는 활성 연도 {activeYears}년에 걸쳐 선별 논문 {papers}편과 인용 {citations}회를 포함합니다. 가장 강한 taxonomy 신호는 {topCategories}이며, 논문 수가 가장 많은 해는 {peakYear}년({peakYearCount}편)입니다.",
                "인용 기준 최상위 논문은 {topPaperCategory} 분류의 \"{topPaper}\"({topPaperYear}, {topPaperCitations}회 인용)입니다. {topKeywords} 같은 키워드는 이 기간이 측정, 기전, 임상 전환, 계산 모델을 어떻게 연결하는지 보여줍니다.",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {
                    "label": "Period Shape",
                    "title": "선택한 기간이 연구 지형을 바꿉니다",
                    "body": "{topCategory}가 {topCategoryCount}편을 차지해, {range}의 가시적 코퍼스는 이 시기에 강했던 taxonomy 축을 중심으로 구성됩니다.",
                    "implication": "시사점: 하나의 기간만 보고 장기 중심축을 단정하기보다 인접 기간과 비교해야 합니다.",
                },
                {
                    "label": "Citation Mass",
                    "title": "인용 집중은 공유 인프라를 드러냅니다",
                    "body": "이 기간은 총 {citations}회 인용을 가지며, 인용 피크는 {peakCitationYear}년 부근입니다. 고인용 논문은 atlas, protocol, cohort, 분석 도구처럼 재사용 가능한 기반을 정의하는 경우가 많습니다.",
                    "implication": "시사점: 지속적 영향력은 다른 하위 분야가 반복해서 쓰는 자원에서 자주 발생합니다.",
                },
                {
                    "label": "Translation",
                    "title": "임상 신호와 기전 신호를 함께 읽어야 합니다",
                    "body": "최상위 논문 \"{topPaper}\"은 {topPaperCategory}에 속하고, 주요 분류는 {topCategories}입니다. 이 조합은 해당 기간이 질환 cohort, 생물학적 기전, 측정 기술, 계산 모델 중 무엇에 의해 움직였는지 보여줍니다.",
                    "implication": "시사점: 좋은 요약은 환자 relevance와 causal 또는 measurement evidence를 함께 연결해야 합니다.",
                },
                {
                    "label": "Methods",
                    "title": "키워드는 기간별 방법론의 축을 보여줍니다",
                    "body": "{topKeywords} 같은 빈도 높은 태그는 선택된 연도를 조직하는 도구, 모델, 연구 setting을 드러냅니다.",
                    "implication": "시사점: 키워드 변화는 인용 수가 충분히 성숙하기 전에도 유용한 조기 신호가 됩니다.",
                },
                {
                    "label": "Open Gaps",
                    "title": "인용 기반 지도에는 전문가 해석이 필요합니다",
                    "body": "요약이 안정적으로 보여도 최신 연구, 과소대표 집단, null result, protocol 차이는 약하게 보일 수 있습니다.",
                    "implication": "시사점: 이 화면은 탐색 layer로 사용하고, 핵심 주장은 full-text review와 재현성 점검으로 보강해야 합니다.",
                },
            ],
        },
        "zh": {
            "timelineTitle": "研究时间线",
            "summary": [
                "在 {range} 期间，brain 语料包含 {papers} 篇入选论文，覆盖 {activeYears} 个活跃年份，总引用为 {citations} 次。最强的 taxonomy 信号是 {topCategories}，论文数量峰值出现在 {peakYear} 年（{peakYearCount} 篇）。",
                "按引用排序的领先论文是 {topPaperCategory} 中的《{topPaper}》（{topPaperYear}，{topPaperCitations} 次引用）。{topKeywords} 等关键词显示该时期如何连接测量、机制、临床转化和计算模型。",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {
                    "label": "Period Shape",
                    "title": "所选时期会改变领域地图",
                    "body": "{topCategory} 包含 {topCategoryCount} 篇论文，因此 {range} 的可见语料主要由该时期最强的 taxonomy 轴支撑。",
                    "implication": "启示：不要只凭一个时间段判断长期中心，应与相邻时期一起比较。",
                },
                {
                    "label": "Citation Mass",
                    "title": "引用集中度揭示共享基础设施",
                    "body": "该时期共有 {citations} 次引用，引用峰值接近 {peakCitationYear} 年。高引用论文往往定义可复用的 atlas、protocol、cohort 或分析工具。",
                    "implication": "启示：持久影响力常来自可被多个子领域反复使用的资源。",
                },
                {
                    "label": "Translation",
                    "title": "临床信号和机制信号需要一起解读",
                    "body": "领先论文《{topPaper}》属于 {topPaperCategory}，主要分类包括 {topCategories}。这种组合说明该时期更受疾病队列、机制、测量还是计算模型驱动。",
                    "implication": "启示：高质量总结应同时连接患者相关性与因果或测量证据。",
                },
                {
                    "label": "Methods",
                    "title": "关键词呈现时期的方法学主干",
                    "body": "{topKeywords} 等高频标签显示了组织所选年份的工具、模型和研究场景。",
                    "implication": "启示：在引用尚未成熟时，关键词变化可以作为早期趋势信号。",
                },
                {
                    "label": "Open Gaps",
                    "title": "引用排序地图仍需要专家解释",
                    "body": "即使时期摘要看起来稳定，最新研究、代表性不足的人群、阴性结果和 protocol 差异也可能被弱化。",
                    "implication": "启示：把此视图作为导航层，再用全文审阅和可重复性检查支撑核心判断。",
                },
            ],
        },
        "ja": {
            "timelineTitle": "研究タイムライン",
            "summary": [
                "{range} の brain コーパスには、{activeYears} の対象年にわたる選定論文 {papers} 本、引用 {citations} 件が含まれます。最も強い taxonomy 信号は {topCategories} で、論文数のピークは {peakYear} 年（{peakYearCount} 本）です。",
                "引用順で最上位の論文は {topPaperCategory} の「{topPaper}」（{topPaperYear}、{topPaperCitations} 件引用）です。{topKeywords} などのキーワードは、この期間が計測、機序、臨床応用、計算モデルをどう結びつけているかを示します。",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {
                    "label": "Period Shape",
                    "title": "選択した期間によって研究地図は変わります",
                    "body": "{topCategory} が {topCategoryCount} 本を占めるため、{range} の可視化されたコーパスはこの時期に強かった taxonomy 軸を中心に構成されます。",
                    "implication": "示唆：単一期間だけで長期的な中心領域を判断せず、隣接期間と比較する必要があります。",
                },
                {
                    "label": "Citation Mass",
                    "title": "引用の集中は共有インフラを示します",
                    "body": "この期間の引用は合計 {citations} 件で、引用ピークは {peakCitationYear} 年付近です。高引用論文は atlas、protocol、cohort、解析ツールなど再利用可能な基盤を定義することがよくあります。",
                    "implication": "示唆：持続的な影響力は、複数の下位分野が再利用できる資源から生まれやすいです。",
                },
                {
                    "label": "Translation",
                    "title": "臨床シグナルと機序シグナルを合わせて読む必要があります",
                    "body": "最上位論文「{topPaper}」は {topPaperCategory} に属し、主要分類は {topCategories} です。この組み合わせは、その期間が疾患 cohort、機序、計測、計算モデルのどれに動かされているかを示します。",
                    "implication": "示唆：良い要約は、患者への関連性と因果または計測エビデンスを結びつける必要があります。",
                },
                {
                    "label": "Methods",
                    "title": "キーワードは期間ごとの方法論の軸を示します",
                    "body": "{topKeywords} などの高頻度タグは、選択された年を支える装置、モデル、研究設定を示します。",
                    "implication": "示唆：引用数が十分に成熟する前でも、キーワードの変化は早期の傾向信号になります。",
                },
                {
                    "label": "Open Gaps",
                    "title": "引用ベースの地図には専門的解釈が必要です",
                    "body": "要約が安定して見えても、最新研究、過小代表集団、null result、protocol の違いは弱く見えることがあります。",
                    "implication": "示唆：この画面はナビゲーション層として使い、主要な判断は全文レビューと再現性確認で補強してください。",
                },
            ],
        },
    }


def overall_period_summary(rows, start, end, by_year):
    counts, citations = category_stats(rows)
    keywords = keyword_stats(rows)
    year_counts = Counter(row["year"] for row in rows)
    year_citations = defaultdict(int)
    for row in rows:
        year_citations[row["year"]] += row["citationCount"]
    peak_year, peak_count = year_counts.most_common(1)[0] if year_counts else (None, 0)
    peak_citation_year = max(year_citations, key=year_citations.get) if year_citations else None
    top = max(rows, key=lambda item: item["citationCount"]) if rows else None
    summary = {
        "startYear": start,
        "endYear": end,
        "rangeLabel": str(start) if start == end else f"{start}-{end}",
        "totalPapers": len(rows),
        "activeYears": sum(1 for year in range(start, end + 1) if by_year[year]),
        "citationCount": sum(row["citationCount"] for row in rows),
        "categoryCount": len(counts),
        "topCategories": [
            {"name": category, "count": count, "citations": citations[category]}
            for category, count in counts.most_common(6)
        ],
        "topKeywords": [
            {"name": keyword.strip(), "count": count}
            for keyword, count in keywords.most_common(6)
        ],
        "peakYear": peak_year,
        "peakYearCount": peak_count,
        "peakCitationYear": peak_citation_year,
        "peakCitationCount": year_citations.get(peak_citation_year, 0) if peak_citation_year else 0,
        "topPaper": {
            "title": top["title"],
            "year": top["year"],
            "category": top["category"],
            "url": top["url"],
            "citations": top["citationCount"],
        } if top else None,
    }
    summary["periodInsights"] = paper_curation_period_insights(summary, "brain")
    return summary


def _insight_number(value):
    return f"{int(value or 0):,}"


def _insight_names(items, key="name", limit=3, fallback="metadata-ranked signals"):
    values = [
        str(item.get(key, "")).strip()
        for item in (items or [])[:limit]
        if str(item.get(key, "")).strip()
    ]
    return ", ".join(values) if values else fallback


def paper_curation_period_insights(summary, corpus_label):
    categories = summary.get("topCategories") or []
    keywords = summary.get("topKeywords") or []
    top_category = categories[0] if categories else {}
    top_paper = summary.get("topPaper") or {}
    range_label = summary.get("rangeLabel") or f"{summary.get('startYear')}-{summary.get('endYear')}"
    top_category_name = top_category.get("name") or "the leading taxonomy"
    top_category_count = _insight_number(top_category.get("count"))
    category_names = _insight_names(categories)
    keyword_names = _insight_names(keywords, fallback="keyword convention signals")
    citations = _insight_number(summary.get("citationCount"))
    papers = _insight_number(summary.get("totalPapers"))
    peak_year = summary.get("peakYear") or "n/a"
    peak_citation_year = summary.get("peakCitationYear") or "n/a"
    top_paper_title = top_paper.get("title") or "the leading paper"
    top_paper_year = top_paper.get("year") or "n/a"
    top_paper_citations = _insight_number(top_paper.get("citations"))
    top_paper_category = top_paper.get("category") or top_category_name

    return {
        "en": [
            {
                "label": "Period Shape",
                "title": f"{range_label} is led by {top_category_name}",
                "body": f"{top_category_name} contributes {top_category_count} selected papers, with {category_names} forming the visible taxonomy backbone for this period.",
                "implication": "Implication: period changes should be read as changes in the research map, not just as a shorter paper list.",
            },
            {
                "label": "Citation Backbone",
                "title": f"{peak_citation_year} carries the citation center of gravity",
                "body": f"The selected {corpus_label} range contains {papers} papers and {citations} citations; citation mass peaks around {peak_citation_year}, while paper volume peaks around {peak_year}.",
                "implication": "Implication: paper-curation period views help separate durable infrastructure papers from fast-moving recent topics.",
            },
            {
                "label": "Keyword Signals",
                "title": f"Keyword evidence points to {keyword_names}",
                "body": f"The period-specific keyword convention highlights {keyword_names}, giving a method-level reading that complements taxonomy and citation rank.",
                "implication": "Implication: keyword shifts are useful early signals before citations fully mature.",
            },
            {
                "label": "Representative Paper",
                "title": f"\"{top_paper_title}\" anchors the selected range",
                "body": f"The leading citation-ranked paper is from {top_paper_year} in {top_paper_category}, with {top_paper_citations} citations.",
                "implication": "Implication: use the top paper as an entry point, then compare it with adjacent taxonomy clusters before drawing broad conclusions.",
            },
            {
                "label": "Review Priority",
                "title": "What deserves full-text review next",
                "body": f"This metadata-adapter insight is generated in the spirit of paper-curation: it flags {category_names} and {keyword_names} as the period's highest-priority reading lanes.",
                "implication": "Implication: full PDF review remains the next step for causal claims, reproducibility, and experimental detail.",
            },
        ]
    }


def write_overall_analysis(selected):
    by_year = {year: [row for row in selected if row["year"] == year] for year in YEARS}
    analysis = {}
    for start in YEARS:
        period_rows = []
        for end in range(start, END_YEAR + 1):
            period_rows.extend(by_year[end])
            analysis[f"{start}-{end}"] = overall_period_summary(period_rows, start, end, by_year)
    payload = {
        "generated": GENERATED_DATE,
        "yearRange": YEAR_RANGE_TEXT,
        "languages": LANGUAGES,
        "uiLabels": UI_LABELS,
        "analysis": analysis,
    }
    (DATA_DIR / OVERALL_ANALYSIS_JSON).write_text(json.dumps(payload, ensure_ascii=False, separators=(",", ":")), encoding="utf-8")


def write_period_analysis(selected):
    by_year = {year: [row for row in selected if row["year"] == year] for year in YEARS}
    analysis = {}
    for start in YEARS:
        period_rows = []
        for end in range(start, END_YEAR + 1):
            period_rows.extend(by_year[end])
            counts, citations = category_stats(period_rows)
            keywords = keyword_stats(period_rows)
            top = max(period_rows, key=lambda item: item["citationCount"]) if period_rows else None
            analysis[f"{start}_{end}"] = {
                "startYear": start,
                "endYear": end,
                "totalPapers": len(period_rows),
                "activeYears": sum(1 for year in range(start, end + 1) if by_year[year]),
                "citationCount": sum(row["citationCount"] for row in period_rows),
                "categoryCounts": dict(counts),
                "categoryCitations": dict(citations),
                "keywordCounts": dict(keywords),
                "topPaper": {"title": top["title"], "year": top["year"], "url": top["url"], "citations": top["citationCount"]} if top else None,
            }
    (DATA_DIR / PERIOD_ANALYSIS_JSON).write_text(json.dumps(analysis, ensure_ascii=False, indent=2), encoding="utf-8")


def badge(keyword):
    color = KEYWORD_COLORS.get(keyword, "64748b")
    label = keyword.replace("-", "--")
    return f"![{keyword}](https://img.shields.io/badge/keyword-{label}-{color})"


def html_badge(keyword):
    color = KEYWORD_COLORS.get(keyword, "64748b")
    return f'<span class="badge" style="--badge-color: #{html.escape(color)}">{html.escape(keyword)}</span>'


def markdown_link(title, url):
    if url:
        return f"[{title}]({url})"
    return title


def html_escape(value):
    return html.escape(clean_text(value), quote=True)


def paper_table(rows, limit=20):
    out = [
        '<table width="100%">',
        "<thead><tr><th align=\"right\">Rank</th><th>Paper</th><th>Meta</th><th>Keywords</th><th>Key idea</th><th>Strengths</th><th>Limitations</th></tr></thead>",
        "<tbody>",
    ]
    for index, row in enumerate(rows[:limit], start=1):
        link = markdown_link(html_escape(row["title"]), html_escape(row["url"]))
        keywords = " ".join(badge(tag) for tag in row["keywordTags"].split(";") if tag) or "-"
        out.append(
            "<tr>"
            f"<td align=\"right\">{index}</td>"
            f"<td>{link}<br><sub>{html_escape(row['authors'])}</sub></td>"
            f"<td>{row['year']}<br>{html_escape(row['venue'])}<br>{row['citationCount']:,} citations</td>"
            f"<td>{keywords}</td>"
            f"<td>{html_escape(row['keyIdea'])}</td>"
            f"<td>{html_escape(row['strengths'])}</td>"
            f"<td>{html_escape(row['limitations'])}</td>"
            "</tr>"
        )
    if len(rows) > limit:
        out.append(f"<tr><td colspan=\"7\">See the website and taxonomy CSV for all {len(rows):,} papers in this category.</td></tr>")
    out.extend(["</tbody>", "</table>"])
    return "\n".join(out)


def markdown_to_html_doc(title, markdown):
    body_lines = []
    in_table = False
    for line in markdown.splitlines():
        if line.startswith("# "):
            body_lines.append(f"<h1>{html_escape(line[2:])}</h1>")
        elif line.startswith("## "):
            body_lines.append(f"<h2>{html_escape(line[3:])}</h2>")
        elif line.startswith("### "):
            body_lines.append(f"<h3>{html_escape(line[4:])}</h3>")
        elif line.startswith("|"):
            if not in_table:
                body_lines.append("<pre class=\"table-block\">")
                in_table = True
            body_lines.append(html_escape(line))
        else:
            if in_table:
                body_lines.append("</pre>")
                in_table = False
            if not line.strip():
                body_lines.append("")
            elif line.startswith("- "):
                body_lines.append(f"<p>{html_escape(line)}</p>")
            else:
                body_lines.append(f"<p>{html_escape(line)}</p>")
    if in_table:
        body_lines.append("</pre>")
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html_escape(title)}</title>
  <style>
    body {{ font-family: Inter, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; max-width: 980px; margin: 40px auto; padding: 0 24px; color: #172033; line-height: 1.65; }}
    h1, h2, h3 {{ line-height: 1.2; }}
    .table-block {{ white-space: pre-wrap; background: #f6f8fb; border: 1px solid #d9dee8; padding: 14px; overflow: auto; }}
  </style>
</head>
<body>
{chr(10).join(body_lines)}
</body>
</html>
"""


def write_readme(selected, candidates):
    counts, citations = category_stats(selected)
    stats = year_stats(selected)
    lines = [
        "# Awesome Brain",
        "",
        "[![Awesome](https://awesome.re/badge-flat.png)](https://awesome.re)",
        "",
        "A taxonomy-first, citation-ranked map of brain research from 1900 through 2026.",
        "",
        '<p align="center">',
        '  <a href="https://honggi82.github.io/awesome-brain/">',
        '    <img src="https://img.shields.io/badge/Open_Interactive_Website-honggi82.github.io%2Fawesome--brain-0f766e?style=for-the-badge" alt="Open Interactive Website">',
        "  </a>",
        "</p>",
        "",
        "> Browse the full interactive taxonomy site with period, keyword, chart, and paper-card filters: https://honggi82.github.io/awesome-brain/",
        "",
        f"Generated on {GENERATED_DATE} from free public Semantic Scholar metadata. The current edition investigates up to {CANDIDATES_PER_YEAR:,} brain-related candidate papers per year for {YEAR_RANGE_TEXT}, keeps an audited candidate pool, selects the top {TARGET_PER_YEAR} papers per year by citation count, and reorganizes the selected {len(selected):,} papers by research taxonomy.",
        "",
        "## Project Links",
        "",
        "- Website: https://honggi82.github.io/awesome-brain/",
        f"- Selected dataset: `data/{PAPERS_CSV}`",
        f"- Taxonomy dataset with paper-level ideas, strengths, and limitations: `data/{TAXONOMY_CSV}`",
        f"- Precomputed period and keyword analysis: `data/{PERIOD_ANALYSIS_JSON}`",
        f"- Candidate pool: `data/{CANDIDATES_CSV}`",
        "- English review draft: `paper/review_en.html`, `paper/review_en.docx`",
        "- Korean review draft: `paper/review_ko.html`",
        "- Curation method: `paper/curation_method.md`, `paper/curation_method.html`",
        "",
        "## Keywords Convention",
        "",
        "These badges define the brain keyword tags used to read and extend this collection.",
        "",
    ]
    for keyword, description, _ in KEYWORD_CONVENTION:
        lines.append(f"- {badge(keyword)} **{keyword}**: {description}")
    lines.extend(["", "## Taxonomy Overview", ""])
    lines.append(f"- **Total selected papers**: {len(selected):,} papers")
    for category, count in counts.most_common():
        lines.append(f"- **{category}**: {count:,} papers")
    lines.extend(["", "## Taxonomy Collections", ""])
    by_category = defaultdict(list)
    for row in selected:
        by_category[row["category"]].append(row)
    for category in CATEGORIES:
        name = category["name"]
        rows = sorted(by_category[name], key=lambda item: (-item["citationCount"], item["year"], item["title"].lower()))
        if not rows:
            continue
        years = [row["year"] for row in rows]
        lines.extend(
            [
                f"### {name}",
                "",
                f"- Papers selected: **{len(rows):,}**",
                f"- Years covered: **{min(years)}-{max(years)}**",
                f"- Citation count in selected set: **{sum(row['citationCount'] for row in rows):,}**",
                "- Category Overview (main research trends):",
            ]
        )
        lines.extend([f"  - {item}" for item in category["overview"]])
        lines.append("- Limitations:")
        lines.extend([f"  - {item}" for item in category["limitations"]])
        lines.extend(
            [
                "",
                f"<details>",
                f"<summary><strong>Show representative papers for {name}</strong></summary>",
                "",
                paper_table(rows),
                "",
                "</details>",
                "",
            ]
        )
    lines.extend(["## Yearly Coverage", ""])
    lines.append("| Year | Candidate papers | Selected papers | Citation count | Top paper |")
    lines.append("| ---: | ---: | ---: | ---: | --- |")
    candidate_counts = Counter(row["year"] for row in candidates)
    for year in YEARS:
        info = stats.get(year)
        if not info:
            lines.append(f"| {year} | {candidate_counts.get(year, 0):,} | 0 | 0 | - |")
            continue
        top = info["top"]
        lines.append(
            f"| {year} | {candidate_counts.get(year, 0):,} | {info['count']:,} | {info['citations']:,} | {markdown_link(top['title'], top['url'])} |"
        )
    lines.extend(
        [
            "",
            "## Method",
            "",
            f"The collection uses the Semantic Scholar Academic Graph bulk paper search. For each publication year from {START_YEAR} through {END_YEAR}, the pipeline runs broad brain and neuroscience queries, requests results sorted by `citationCount:desc`, merges and deduplicates the results, keeps up to {CANDIDATES_PER_YEAR:,} brain-relevant candidates per year after local title/abstract/field relevance checks, and selects the top {TARGET_PER_YEAR} papers per year by citation count. Taxonomy, keyword tags, key ideas, strengths, limitations, and audit scores are deterministic rule-based enrichments so the repository can be regenerated without paid APIs.",
            "",
            "## Caveats",
            "",
            "- Citation counts favor older papers and can under-rank very recent 2026 work.",
            "- Some early years have fewer than 1,000 discoverable candidates because the public metadata pool is sparse before large-scale journal indexing; every year still contributes 100 selected papers.",
            "- Semantic Scholar metadata is broad scholarly metadata; this is not a full systematic review of every PDF.",
            "- Influential citation counts use Semantic Scholar metadata when available.",
            "- Very broad brain research spans neuroscience, neurology, psychology, molecular biology, computation, and clinical medicine, so taxonomy boundaries are necessarily approximate.",
            "",
            "## Acknowledgements",
            "",
            "This repository and interactive site were created with appreciation for [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation). Its paper-curation workflow and repository organization informed the approach used here for a taxonomy-first, citation-ranked research map.",
        ]
    )
    readme = "\n".join(lines) + "\n"
    (ROOT / "README.md").write_text(readme, encoding="utf-8")
    (ROOT / "README.html").write_text(markdown_to_html_doc("Awesome Brain", readme), encoding="utf-8")


def write_taxonomy_assets():
    target = DOCS_DIR / "assets" / "taxonomy"
    missing = []
    for category in CATEGORIES:
        image_path = target / f"{category['slug']}.png"
        if not image_path.exists():
            missing.append(image_path.as_posix())
    if missing:
        raise FileNotFoundError("Missing taxonomy PNG assets: " + ", ".join(missing))


def site_rows(selected):
    rows = []
    for row in selected:
        rows.append(
            {
                "rank": row["rank"],
                "year": row["year"],
                "title": row["title"],
                "authors": row["authors"],
                "venue": row["venue"],
                "citationCount": row["citationCount"],
                "influentialCitationCount": row["influentialCitationCount"],
                "importanceScore": row["importanceScore"],
                "category": row["category"],
                "keywordTags": [tag for tag in row["keywordTags"].split(";") if tag],
                "keyIdea": row["keyIdea"],
                "strengths": row["strengths"],
                "limitations": row["limitations"],
                "url": row["url"],
                "semanticScholarUrl": semantic_scholar_url(row),
                "openAccessPdf": row["openAccessPdf"],
                "githubUrl": row.get("githubUrl", ""),
                "workType": row["workType"],
            }
        )
    return rows


def write_site(selected):
    apply_github_links(selected)
    categories = [
        {
            "name": item["name"],
            "slug": item["slug"],
            "accent": item["accent"],
            "overview": item["overview"],
            "limitations": item["limitations"],
            "icon": f"assets/taxonomy/{item['slug']}.png",
        }
        for item in CATEGORIES
    ]
    keyword_info = [{"name": keyword, "description": desc, "color": color} for keyword, desc, color in KEYWORD_CONVENTION]
    payload = json.dumps(site_rows(selected), ensure_ascii=False)
    category_payload = json.dumps(categories, ensure_ascii=False)
    keyword_payload = json.dumps(keyword_info, ensure_ascii=False)
    research_overview = research_overview_html().strip()
    research_copy_payload = json.dumps(research_copy(), ensure_ascii=False)
    overall_research_templates_payload = json.dumps(overall_research_templates(), ensure_ascii=False)
    ui_labels_payload = json.dumps(UI_LABELS, ensure_ascii=False)
    language_options = "\n".join(
        f'<option value="{code}"{" selected" if code == "en" else ""}>{html.escape(label)}</option>'
        for code, label in LANGUAGES.items()
    )
    index = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Awesome Brain</title>
  <style>
    :root {{
      color-scheme: light;
      --ink: #142033;
      --muted: #5b6678;
      --line: #dbe3ee;
      --soft: #f6f8fb;
      --panel: #ffffff;
      --accent: #2563eb;
      --focus: #0f766e;
    }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; color: var(--ink); background: #ffffff; }}
    header {{ padding: 26px 28px 18px; border-bottom: 1px solid var(--line); background: linear-gradient(180deg, #f8fafc 0%, #ffffff 100%); }}
    h1 {{ margin: 0 0 8px; font-size: 34px; letter-spacing: 0; }}
    h2 {{ margin: 0 0 12px; font-size: 22px; }}
    h3 {{ margin: 0; font-size: 18px; }}
    p {{ color: var(--muted); line-height: 1.55; }}
    a {{ color: #1d4ed8; }}
    .topline {{ display: flex; flex-wrap: wrap; gap: 12px; align-items: center; justify-content: space-between; }}
    .links {{ display: flex; flex-wrap: wrap; gap: 10px; }}
    .links a, button, select {{ min-height: 38px; border: 1px solid var(--line); background: var(--panel); color: var(--ink); border-radius: 6px; padding: 8px 11px; font: inherit; }}
    button {{ cursor: pointer; }}
    button:hover, select:hover {{ border-color: #94a3b8; }}
    main {{ padding: 24px 28px 48px; max-width: 1440px; margin: 0 auto; }}
    .controls {{ display: grid; grid-template-columns: repeat(6, minmax(130px, 1fr)); gap: 12px; align-items: end; padding: 14px 0 18px; border-bottom: 1px solid var(--line); }}
    label {{ display: grid; gap: 6px; font-size: 12px; color: var(--muted); font-weight: 700; text-transform: uppercase; }}
    .stats {{ display: grid; grid-template-columns: repeat(4, minmax(160px, 1fr)); gap: 12px; margin: 18px 0; }}
    .stat {{ border: 1px solid var(--line); border-radius: 8px; padding: 14px; background: var(--soft); }}
    .stat strong {{ display: block; font-size: 26px; color: var(--ink); }}
    .research-brief {{ margin: 28px 0; padding: 24px 0; border-top: 1px solid var(--line); border-bottom: 1px solid var(--line); }}
    .timeline-copy {{ max-width: 1080px; }}
    .research-insights {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(260px, 1fr)); gap: 12px; margin-top: 12px; }}
    .insight-box {{ border: 1px solid var(--line); border-radius: 8px; padding: 14px; background: #fff; }}
    .insight-label {{ color: var(--focus); font-size: 12px; font-weight: 800; text-transform: uppercase; letter-spacing: 0; }}
    .insight-box h3 {{ margin: 6px 0 8px; font-size: 17px; }}
    .insight-box p {{ margin: 8px 0 0; }}
    .insight-implication {{ color: var(--ink); font-weight: 700; }}
    .keyword-panel {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(240px, 1fr)); gap: 10px; margin: 12px 0 8px; }}
    .keyword-button {{ display: grid; gap: 8px; align-items: start; min-height: 112px; border-radius: 8px; text-align: left; line-height: 1.45; }}
    .keyword-button[aria-pressed="true"] {{ border-color: var(--focus); box-shadow: 0 0 0 2px rgba(15, 118, 110, 0.16); }}
    .keyword-chip {{ justify-self: start; display: inline-flex; gap: 7px; align-items: center; border-radius: 999px; padding: 4px 9px; background: var(--keyword-color); color: white; font-size: 13px; font-weight: 800; }}
    .keyword-dot {{ width: 8px; height: 8px; border-radius: 50%; background: currentColor; display: inline-block; opacity: 0.88; }}
    .keyword-description {{ display: block; width: 100%; color: var(--muted); }}
    .keyword-count {{ color: var(--focus); font-size: 12px; font-weight: 800; }}
    .keyword-status {{ min-height: 24px; color: var(--muted); margin: 4px 0 18px; }}
    .charts {{ display: grid; grid-template-columns: minmax(260px, 1fr) minmax(260px, 1fr); gap: 18px; margin: 18px 0 24px; }}
    .chart-wrap {{ border: 1px solid var(--line); border-radius: 8px; padding: 14px; min-height: 280px; }}
    canvas {{ width: 100%; height: 230px; display: block; }}
    .taxonomy {{ display: grid; gap: 12px; }}
    details {{ border: 1px solid var(--line); border-radius: 8px; background: var(--panel); overflow: hidden; }}
    summary {{ list-style: none; cursor: pointer; padding: 0; }}
    summary::-webkit-details-marker {{ display: none; }}
    .summary-row {{ display: grid; grid-template-columns: 98px minmax(220px, 1fr) repeat(3, minmax(110px, 150px)); gap: 14px; align-items: center; padding: 12px 14px; }}
    .summary-row img, .summary-all-icon {{ width: 98px; height: 56px; object-fit: cover; border-radius: 6px; border: 1px solid var(--line); }}
    .summary-all-icon {{ display: inline-flex; align-items: center; justify-content: center; background: #eef6ff; color: #1d4ed8; font-weight: 800; }}
    .summary-metric {{ color: var(--muted); font-size: 13px; }}
    .summary-metric strong {{ display: block; color: var(--ink); font-size: 18px; }}
    .section-body {{ padding: 16px; border-top: 1px solid var(--line); }}
    .section-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 18px; margin-bottom: 16px; }}
    .section-grid ul {{ margin-top: 8px; color: var(--muted); line-height: 1.55; }}
    .paper-list {{ display: grid; gap: 10px; }}
    .paper-card {{ border: 1px solid var(--line); border-radius: 8px; padding: 13px; background: #fff; }}
    .paper-head {{ display: flex; gap: 10px; justify-content: space-between; align-items: start; }}
    .paper-title {{ font-weight: 800; line-height: 1.35; }}
    .meta {{ color: var(--muted); font-size: 13px; margin-top: 4px; }}
    .badges {{ display: flex; flex-wrap: wrap; gap: 6px; margin: 10px 0; }}
    .badge {{ display: inline-flex; align-items: center; border: 1px solid color-mix(in srgb, var(--badge-color), #ffffff 35%); background: color-mix(in srgb, var(--badge-color), #ffffff 88%); color: #172033; border-radius: 999px; padding: 3px 8px; font-size: 12px; }}
    .paper-card dl {{ display: grid; grid-template-columns: 110px 1fr; gap: 7px 12px; margin: 10px 0 0; }}
    .paper-card dt {{ font-weight: 800; color: var(--ink); }}
    .paper-card dd {{ margin: 0; color: var(--muted); }}
    .empty {{ padding: 22px; border: 1px dashed var(--line); border-radius: 8px; color: var(--muted); }}
    @media (max-width: 920px) {{
      .controls, .stats, .charts, .section-grid {{ grid-template-columns: 1fr; }}
      .summary-row {{ grid-template-columns: 74px 1fr; }}
      .summary-metric {{ grid-column: 2; }}
      .summary-row img, .summary-all-icon {{ width: 74px; height: 52px; }}
      .paper-head, .paper-card dl {{ display: block; }}
      .paper-card dt {{ margin-top: 8px; }}
    }}
  </style>
</head>
<body>
  <header>
    <div class="topline">
      <div>
        <h1>Awesome Brain</h1>
        <p>Taxonomy-first, citation-ranked brain research map for {YEAR_RANGE_TEXT}. Generated {GENERATED_DATE} from free Semantic Scholar metadata.</p>
      </div>
      <nav class="links" aria-label="Project links">
        <a href="../README.md">README</a>
        <a href="../data/{PAPERS_CSV}">Selected CSV</a>
        <a href="../data/{CANDIDATES_CSV}">Candidate CSV</a>
        <a href="../paper/review_en.html">Review</a>
      </nav>
    </div>
  </header>
  <main>
    <section class="controls" aria-label="Filters">
      <label>Preset
        <select id="preset">
          <option value="{START_YEAR}-{END_YEAR}">All years</option>
          <option value="1900-1949">1900-1949</option>
          <option value="1950-1979">1950-1979</option>
          <option value="1980-1999">1980-1999</option>
          <option value="2000-2009">2000-2009</option>
          <option value="2010-2019">2010-2019</option>
          <option value="2020-2026">2020-2026</option>
        </select>
      </label>
      <label>Start year<select id="startYear"></select></label>
      <label>End year<select id="endYear"></select></label>
      <label>Language
        <select id="language">
          {language_options}
        </select>
      </label>
      <button id="reset" type="button">Reset</button>
    </section>
    <section class="stats" id="stats" aria-live="polite"></section>
    {research_overview}
    <h2>Keywords Convention</h2>
    <div class="keyword-panel" id="keywordPanel"></div>
    <div class="keyword-status" id="keywordStatus">No keyword selected.</div>
    <section class="charts">
      <div class="chart-wrap"><h2>Category Distribution</h2><canvas id="categoryChart" width="720" height="300"></canvas></div>
      <div class="chart-wrap"><h2>Yearly Citation Mass</h2><canvas id="yearChart" width="720" height="300"></canvas></div>
    </section>
    <section class="taxonomy" id="taxonomy"></section>
  </main>
  <script>
    const PAPERS = {payload};
    const CATEGORIES = {category_payload};
    const KEYWORDS = {keyword_payload};
    const RESEARCH_COPY = {research_copy_payload};
    const OVERALL_RESEARCH_TEMPLATES = {overall_research_templates_payload};
    const OVERALL_ANALYSIS_URL = 'data/{OVERALL_ANALYSIS_JSON}';
    const START_YEAR = {START_YEAR};
    const END_YEAR = {END_YEAR};
    const labels = {ui_labels_payload};
    const state = {{ start: START_YEAR, end: END_YEAR, keyword: null, lang: 'en' }};
    let overallAnalysis = null;
    const startSelect = document.getElementById('startYear');
    const endSelect = document.getElementById('endYear');
    const presetSelect = document.getElementById('preset');
    const languageSelect = document.getElementById('language');
    const keywordPanel = document.getElementById('keywordPanel');
    const keywordStatus = document.getElementById('keywordStatus');
    const taxonomy = document.getElementById('taxonomy');

    function fmt(value) {{ return Number(value || 0).toLocaleString(); }}
    function filtered() {{
      return PAPERS.filter(p => p.year >= state.start && p.year <= state.end && (!state.keyword || p.keywordTags.includes(state.keyword)));
    }}
    function fillYears() {{
      for (let year = START_YEAR; year <= END_YEAR; year++) {{
        startSelect.append(new Option(year, year));
        endSelect.append(new Option(year, year));
      }}
      startSelect.value = state.start;
      endSelect.value = state.end;
    }}
    function renderStats(rows) {{
      const activeYears = new Set(rows.map(p => p.year)).size;
      const citations = rows.reduce((sum, p) => sum + p.citationCount, 0);
      const cats = new Set(rows.map(p => p.category)).size;
      const l = labels[state.lang];
      document.getElementById('stats').innerHTML = `
        <div class="stat"><strong>${{fmt(rows.length)}}</strong>${{l.papers}}</div>
        <div class="stat"><strong>${{fmt(activeYears)}}</strong>${{l.years}}</div>
        <div class="stat"><strong>${{fmt(citations)}}</strong>${{l.citations}}</div>
        <div class="stat"><strong>${{fmt(cats)}}</strong>${{l.categories}}</div>`;
    }}
    function escapeHtml(value) {{
      const escapeMap = {{ '&': '&amp;', '<': '&lt;', '>': '&gt;', '"': '&quot;', "'": '&#39;' }};
      return String(value ?? '').replace(/[&<>"']/g, ch => escapeMap[ch]);
    }}
    function names(items, key = 'name') {{
      return (items || []).slice(0, 3).map(item => item[key]).filter(Boolean).join(', ') || 'n/a';
    }}
    function researchTemplateData(metric) {{
      const topCategories = metric.topCategories || [];
      const topKeywords = metric.topKeywords || [];
      const topCategory = topCategories[0] || {{}};
      const topPaper = metric.topPaper || {{}};
      return {{
        range: metric.rangeLabel || `${{metric.startYear}}-${{metric.endYear}}`,
        papers: fmt(metric.totalPapers),
        activeYears: fmt(metric.activeYears),
        citations: fmt(metric.citationCount),
        topCategories: names(topCategories),
        topCategory: topCategory.name || 'n/a',
        topCategoryCount: fmt(topCategory.count || 0),
        topKeywords: names(topKeywords),
        peakYear: metric.peakYear || 'n/a',
        peakYearCount: fmt(metric.peakYearCount || 0),
        peakCitationYear: metric.peakCitationYear || 'n/a',
        topPaper: topPaper.title || 'n/a',
        topPaperYear: topPaper.year || 'n/a',
        topPaperCategory: topPaper.category || 'n/a',
        topPaperCitations: fmt(topPaper.citations || 0)
      }};
    }}
    function applyTemplate(template, data) {{
      let output = template || '';
      Object.keys(data).forEach(key => {{
        output = output.split('{{' + key + '}}').join(escapeHtml(data[key]));
      }});
      return output;
    }}
    function renderOverallResearch(metric) {{
      const copy = OVERALL_RESEARCH_TEMPLATES[state.lang] || OVERALL_RESEARCH_TEMPLATES.en;
      const data = researchTemplateData(metric);
      const summaryHtml = (copy.summary || []).map(text => `<p>${{applyTemplate(text, data)}}</p>`).join('');
      const insightItems = metric?.periodInsights?.[state.lang] || metric?.periodInsights?.en || copy.insights || [];
      const insightHtml = insightItems.map(item => `
        <article class="insight-box">
          <div class="insight-label">${{escapeHtml(item.label)}}</div>
          <h3>${{applyTemplate(item.title, data)}}</h3>
          <p>${{applyTemplate(item.body, data)}}</p>
          <p class="insight-implication">${{applyTemplate(item.implication, data)}}</p>
        </article>`).join('');
      return `
        <h2 id="research-timeline-title">${{escapeHtml(copy.timelineTitle)}}</h2>
        <div class="timeline-copy">${{summaryHtml}}</div>
        <h2>${{escapeHtml(copy.insightsTitle)}}</h2>
        <div class="research-insights">${{insightHtml}}</div>`;
    }}
    function renderResearchCopy() {{
      const brief = document.getElementById('researchBrief');
      if (!brief) return;
      const key = `${{state.start}}-${{state.end}}`;
      const metric = overallAnalysis?.analysis?.[key];
      brief.innerHTML = metric ? renderOverallResearch(metric) : (RESEARCH_COPY[state.lang] || RESEARCH_COPY.en);
    }}
    function renderKeywords(rows) {{
      keywordPanel.innerHTML = KEYWORDS.map(k => {{
        const matchCount = PAPERS.filter(p => p.year >= state.start && p.year <= state.end && p.keywordTags.includes(k.name)).length;
        const pressed = state.keyword === k.name ? 'true' : 'false';
        return `<button class="keyword-button" type="button" data-keyword="${{escapeHtml(k.name)}}" aria-pressed="${{pressed}}" style="--keyword-color:#${{k.color}}">
          <span class="keyword-chip"><span class="keyword-dot"></span><span>${{escapeHtml(k.name)}}</span></span>
          <span class="keyword-description">${{escapeHtml(k.description)}}</span>
          <span class="keyword-count">${{fmt(matchCount)}} papers</span>
        </button>`;
      }}).join('');
      if (state.keyword) {{
        keywordStatus.textContent = `${{state.keyword}} selected - ${{fmt(rows.length)}} matching papers.`;
      }} else {{
        keywordStatus.textContent = labels[state.lang].noKeyword;
      }}
    }}
    keywordPanel.addEventListener('click', event => {{
      const button = event.target.closest('[data-keyword]');
      if (!button) return;
      const keyword = button.dataset.keyword;
      state.keyword = state.keyword === keyword ? null : keyword;
      render();
    }});
    function countsBy(rows, key) {{
      const counts = new Map();
      for (const row of rows) counts.set(row[key], (counts.get(row[key]) || 0) + 1);
      return counts;
    }}
    function citationsByYear(rows) {{
      const counts = new Map();
      for (const row of rows) counts.set(row.year, (counts.get(row.year) || 0) + row.citationCount);
      return counts;
    }}
    function drawBarChart(canvas, entries, color) {{
      const ctx = canvas.getContext('2d');
      const width = canvas.width;
      const height = canvas.height;
      ctx.clearRect(0, 0, width, height);
      ctx.fillStyle = '#ffffff';
      ctx.fillRect(0, 0, width, height);
      const max = Math.max(1, ...entries.map(e => e.value));
      const left = 150, right = 20, top = 16, bottom = 30;
      const plotWidth = width - left - right;
      const rowHeight = Math.max(14, (height - top - bottom) / Math.max(entries.length, 1));
      ctx.font = '12px system-ui';
      entries.forEach((entry, index) => {{
        const y = top + index * rowHeight;
        const barWidth = Math.max(1, (entry.value / max) * plotWidth);
        ctx.fillStyle = '#5b6678';
        ctx.fillText(entry.label.slice(0, 22), 8, y + rowHeight * 0.62);
        ctx.fillStyle = color;
        ctx.fillRect(left, y + 3, barWidth, Math.max(7, rowHeight - 7));
        ctx.fillStyle = '#172033';
        ctx.fillText(fmt(entry.value), left + barWidth + 5, y + rowHeight * 0.62);
      }});
    }}
    function renderCharts(rows) {{
      const catEntries = Array.from(countsBy(rows, 'category'), ([label, value]) => ({{ label, value }})).sort((a, b) => b.value - a.value).slice(0, 10);
      drawBarChart(document.getElementById('categoryChart'), catEntries, '#2563eb');
      const yearCounts = citationsByYear(rows);
      const yearEntries = [];
      for (let year = state.start; year <= state.end; year++) {{
        if (yearCounts.has(year)) yearEntries.push({{ label: String(year), value: yearCounts.get(year) }});
      }}
      const sampled = yearEntries.length > 28 ? yearEntries.filter((_, i) => i % Math.ceil(yearEntries.length / 28) === 0) : yearEntries;
      drawBarChart(document.getElementById('yearChart'), sampled, '#0f766e');
    }}
    function badges(tags) {{
      return tags.map(tag => {{
        const meta = KEYWORDS.find(k => k.name === tag);
        const color = meta ? meta.color : '64748b';
        return `<span class="badge" style="--badge-color:#${{color}}">${{tag}}</span>`;
      }}).join('');
    }}
    function paperCard(p) {{
      const l = labels[state.lang];
      const links = `<a href="${{p.url}}">paper</a>${{p.semanticScholarUrl ? ` · <a href="${{p.semanticScholarUrl}}">Semantic Scholar</a>` : ''}}${{p.openAccessPdf ? ` · <a href="${{p.openAccessPdf}}">PDF</a>` : ''}}${{p.githubUrl ? ` · <a href="${{p.githubUrl}}">GitHub</a>` : ''}}`;
      return `<article class="paper-card" data-year="${{p.year}}" data-keywords="${{p.keywordTags.join(' ')}}">
        <div class="paper-head"><div><a class="paper-title" href="${{p.url}}">${{p.title}}</a><div class="meta">${{p.authors}}</div></div><div class="meta">#${{p.rank}} in ${{p.year}}</div></div>
        <div class="meta">${{p.year}} · ${{p.venue}} · ${{fmt(p.citationCount)}} citations · influential citations ${{fmt(p.influentialCitationCount)}} · score ${{p.importanceScore}} · ${{links}}</div>
        <div class="badges">${{badges(p.keywordTags)}}</div>
        <dl><dt>${{l.keyIdea}}</dt><dd>${{p.keyIdea}}</dd><dt>${{l.strengths}}</dt><dd>${{p.strengths}}</dd><dt>${{l.limitations}}</dt><dd>${{p.limitations}}</dd></dl>
      </article>`;
    }}
    function allTaxonomiesDetails(rows) {{
      const allRows = [...rows].sort((a, b) => b.citationCount - a.citationCount);
      const years = allRows.map(p => p.year);
      const citations = allRows.reduce((sum, p) => sum + p.citationCount, 0);
      const top = allRows[0];
      const visibleCards = allRows.slice(0, 120).map(paperCard).join('');
      const extra = allRows.length > 120 ? `<p class="meta">Showing 120 of ${{fmt(allRows.length)}} matching papers across all taxonomies. The complete set is in the data files.</p>` : '';
      return `<details>
        <summary><div class="summary-row">
          <div class="summary-all-icon" aria-hidden="true">All</div>
          <div><h3>All Taxonomies</h3><div class="meta">Top paper: <a href="${{top.url}}">${{top.title}}</a></div></div>
          <div class="summary-metric"><strong>${{fmt(allRows.length)}}</strong>papers</div>
          <div class="summary-metric"><strong>${{Math.min(...years)}}-${{Math.max(...years)}}</strong>years</div>
          <div class="summary-metric"><strong>${{fmt(citations)}}</strong>citations</div>
        </div></summary>
        <div class="section-body">
          <div class="section-grid">
            <div><h3>Category Overview</h3><p>All papers matching the current period and keyword filters, sorted by citation count across every taxonomy.</p></div>
            <div><h3>Research Limitations</h3><p>Use the individual taxonomy rows below for category-specific context and limitations.</p></div>
          </div>
          <div class="paper-list">${{visibleCards}}</div>${{extra}}
        </div>
      </details>`;
    }}
    function renderTaxonomy(rows) {{
      if (!rows.length) {{
        taxonomy.innerHTML = '<div class="empty">No papers match the current filters.</div>';
        return;
      }}
      const byCategory = new Map();
      for (const category of CATEGORIES) byCategory.set(category.name, []);
      for (const row of rows) byCategory.get(row.category)?.push(row);
      const categoryDetails = CATEGORIES.map(category => {{
        const catRows = (byCategory.get(category.name) || []).sort((a, b) => b.citationCount - a.citationCount);
        if (!catRows.length) return '';
        const years = catRows.map(p => p.year);
        const citations = catRows.reduce((sum, p) => sum + p.citationCount, 0);
        const top = catRows[0];
        const visibleCards = catRows.slice(0, 120).map(paperCard).join('');
        const extra = catRows.length > 120 ? `<p class="meta">Showing 120 of ${{fmt(catRows.length)}} matching papers. The complete set is in the data files.</p>` : '';
        return `<details>
          <summary><div class="summary-row">
            <img src="${{category.icon}}" alt="">
            <div><h3>${{category.name}}</h3><div class="meta">Top paper: <a href="${{top.url}}">${{top.title}}</a></div></div>
            <div class="summary-metric"><strong>${{fmt(catRows.length)}}</strong>papers</div>
            <div class="summary-metric"><strong>${{Math.min(...years)}}-${{Math.max(...years)}}</strong>years</div>
            <div class="summary-metric"><strong>${{fmt(citations)}}</strong>citations</div>
          </div></summary>
          <div class="section-body">
            <div class="section-grid">
              <div><h3>Category Overview</h3><ul>${{category.overview.map(x => `<li>${{x}}</li>`).join('')}}</ul></div>
              <div><h3>Research Limitations</h3><ul>${{category.limitations.map(x => `<li>${{x}}</li>`).join('')}}</ul></div>
            </div>
            <div class="paper-list">${{visibleCards}}</div>${{extra}}
          </div>
        </details>`;
      }}).join('');
      taxonomy.innerHTML = allTaxonomiesDetails(rows) + categoryDetails;
    }}
    function render() {{
      if (state.start > state.end) [state.start, state.end] = [state.end, state.start];
      startSelect.value = state.start;
      endSelect.value = state.end;
      const rows = filtered();
      renderStats(rows);
      renderResearchCopy();
      renderKeywords(rows);
      renderCharts(rows);
      renderTaxonomy(rows);
    }}
    presetSelect.addEventListener('change', () => {{
      const [start, end] = presetSelect.value.split('-').map(Number);
      state.start = start; state.end = end; render();
    }});
    startSelect.addEventListener('change', () => {{ state.start = Number(startSelect.value); render(); }});
    endSelect.addEventListener('change', () => {{ state.end = Number(endSelect.value); render(); }});
    languageSelect.addEventListener('change', () => {{ state.lang = languageSelect.value; render(); }});
    document.getElementById('reset').addEventListener('click', () => {{
      state.start = START_YEAR; state.end = END_YEAR; state.keyword = null; state.lang = 'en';
      presetSelect.value = `${{START_YEAR}}-${{END_YEAR}}`;
      languageSelect.value = 'en';
      render();
    }});
    fillYears();
    render();
    fetch(OVERALL_ANALYSIS_URL)
      .then(response => response.json())
      .then(data => {{ overallAnalysis = data; render(); }})
      .catch(() => {{}});
  </script>
</body>
</html>
"""
    (DOCS_DIR / "index.html").write_text(index, encoding="utf-8")


def reference_line(row):
    return f"{row['authors']}. ({row['year']}). {row['title']}. {row['venue']}. {row['url']}"


def review_sections(selected, korean=False):
    counts, _ = category_stats(selected)
    stats = year_stats(selected)
    total_citations = sum(row["citationCount"] for row in selected)
    leading_category, leading_count = counts.most_common(1)[0]
    peak_year = max(stats, key=lambda year: stats[year]["citations"]) if stats else START_YEAR
    top_cited = sorted(selected, key=lambda row: row["citationCount"], reverse=True)[:15]
    if korean:
        title = f"{YEAR_RANGE_TEXT} brain 연구 동향: 공개 메타데이터 기반 인용순 큐레이션"
        abstract = f"이 리뷰 초안은 Semantic Scholar 공개 메타데이터에서 {START_YEAR}년부터 {END_YEAR}년까지 brain 관련 후보 논문을 연도별 최대 {CANDIDATES_PER_YEAR:,}편 조사하고, 각 연도 인용 상위 {TARGET_PER_YEAR}편을 선정한 결과를 요약한다. 최종 목록은 {len(selected):,}편이며, 신경영상, 전기생리, 세포/분자 신경과학, 인지/시스템 신경과학, 임상 신경학, 발달/연결체, 신경기술, 계산 신경과학, 뇌혈관/손상, 일반 리뷰 분류로 정리했다."
        findings = [
            f"선정 논문은 총 {total_citations:,}회의 인용을 포함하며, 선택 집합에서 인용량이 가장 큰 연도는 {peak_year}년이다.",
            f"가장 큰 분류는 {leading_category}({leading_count:,}편)이다.",
            "MRI, fMRI, EEG, MEG, ECoG, single-cell, human, non-human 등의 키워드 태그는 방법과 대상 축을 빠르게 필터링하기 위해 부여했다.",
            "이 결과는 PDF 전문 심사나 체계적 문헌고찰이 아니라 메타데이터 기반 지도이므로, 해석에는 후속 전문가 검토가 필요하다.",
        ]
        future = [
            "선정 논문의 전문을 읽고 연구 설계, 표본 수, 재현성, 공개 데이터 여부를 별도 코딩한다.",
            "임상 중개, 장기 안정성, 안전성, 윤리, 개인정보 보호 항목을 확장 taxonomy로 추가한다.",
            "최근 연도 논문은 인용 축적 시간이 짧으므로 전문가 추천이나 최신성 보정 점수를 병행한다.",
        ]
        labels = {
            "abstract": "초록",
            "findings": "핵심 발견",
            "taxonomy": "분류별 해석",
            "future": "향후 연구 과제",
            "refs": "선정 참고문헌",
            "top": "인용 상위 논문",
        }
    else:
        title = f"Brain Research from {START_YEAR} to {END_YEAR}: A Metadata-Driven Review"
        abstract = f"This draft review maps brain research from {START_YEAR} through {END_YEAR}, investigating up to {CANDIDATES_PER_YEAR:,} candidate papers per year from free Semantic Scholar metadata and selecting {TARGET_PER_YEAR} papers per year by citation count. The final {len(selected):,} selected papers are organized into neuroimaging, electrophysiology, cellular and molecular neuroscience, cognitive and systems neuroscience, clinical neurology, development and connectomics, neurotechnology, computational neuroscience, cerebrovascular and injury research, and broad reviews."
        findings = [
            f"The selected papers account for {total_citations:,} citations in the selected set, with the largest citation mass in {peak_year}.",
            f"The largest taxonomy category is {leading_category} ({leading_count:,} papers).",
            "Keyword filters expose method and population axes such as MRI, fMRI, EEG, MEG, ECoG, single-cell, human, and non-human.",
            "The result is a metadata-driven map rather than a PDF-level systematic review, so expert appraisal remains necessary before strong field-level claims.",
        ]
        future = [
            "Add full-text appraisal for study design, sample size, reproducibility, data availability, and code release.",
            "Extend the taxonomy with clinical translation, long-term stability, safety, neuroethics, and privacy criteria.",
            "Use expert review or recency-aware scoring to compensate for structurally low citation counts in recent years.",
        ]
        labels = {
            "abstract": "Abstract",
            "findings": "Key Findings",
            "taxonomy": "Category-Level Interpretation",
            "future": "Future Research Agenda",
            "refs": "Selected References",
            "top": "Top Papers by Citation Count",
        }
    category_lines = [f"{cat}: {count:,} papers ({count / len(selected):.1%})" for cat, count in counts.most_common()]
    refs = [reference_line(row) for row in top_cited]
    return title, abstract, findings, category_lines, future, refs, labels, top_cited


def html_ranked_table(rows):
    out = ["<table>", "<thead><tr><th>Year</th><th>Rank</th><th>Paper</th><th>Citations</th><th>Category</th></tr></thead>", "<tbody>"]
    for row in rows:
        out.append(
            f"<tr><td>{row['year']}</td><td>{row['rank']}</td><td><a href=\"{html_escape(row['url'])}\">{html_escape(row['title'])}</a></td><td>{row['citationCount']:,}</td><td>{html_escape(row['category'])}</td></tr>"
        )
    out.extend(["</tbody>", "</table>"])
    return "\n".join(out)


def write_review_html(selected, korean=False):
    title, abstract, findings, category_lines, future, refs, labels, top_cited = review_sections(selected, korean=korean)
    lang = "ko" if korean else "en"
    doc = f"""<!doctype html>
<html lang="{lang}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html_escape(title)}</title>
  <style>
    body {{ font-family: Georgia, "Noto Serif KR", serif; max-width: 920px; margin: 40px auto; padding: 0 22px; line-height: 1.72; color:#172033; }}
    h1 {{ line-height:1.15; }}
    h2 {{ margin-top: 34px; }}
    li {{ margin: 6px 0; }}
    .abstract {{ background:#f5f7fb; border-left:4px solid #2563eb; padding:14px 18px; }}
    table {{ width:100%; border-collapse:collapse; margin:16px 0; }}
    th,td {{ border-bottom:1px solid #d9dee8; padding:8px; vertical-align:top; text-align:left; }}
    th {{ background:#f4f6fa; }}
  </style>
</head>
<body>
  <h1>{html_escape(title)}</h1>
  <p><strong>Generated:</strong> {GENERATED_DATE} &middot; <strong>Dataset:</strong> {len(selected):,} papers</p>
  <h2>{labels['abstract']}</h2>
  <p class="abstract">{html_escape(abstract)}</p>
  <h2>{labels['findings']}</h2>
  <ul>{''.join(f'<li>{html_escape(item)}</li>' for item in findings)}</ul>
  <h2>{labels['taxonomy']}</h2>
  <ul>{''.join(f'<li>{html_escape(item)}</li>' for item in category_lines)}</ul>
  <h2>{labels['top']}</h2>
  {html_ranked_table(top_cited)}
  <h2>{labels['future']}</h2>
  <ul>{''.join(f'<li>{html_escape(item)}</li>' for item in future)}</ul>
  <h2>{labels['refs']}</h2>
  <ol>{''.join(f'<li>{html_escape(ref)}</li>' for ref in refs)}</ol>
</body>
</html>
"""
    filename = "review_ko.html" if korean else "review_en.html"
    (PAPER_DIR / filename).write_text(doc, encoding="utf-8")


def write_review_docx(selected):
    if Document is None:
        (PAPER_DIR / "review_en.docx.txt").write_text("python-docx is not installed; review_en.docx was not generated.\n", encoding="utf-8")
        return
    title, abstract, findings, category_lines, future, refs, labels, top_cited = review_sections(selected, korean=False)
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(f"Generated: {GENERATED_DATE} | Dataset: {len(selected):,} papers")
    document.add_heading(labels["abstract"], level=1)
    document.add_paragraph(abstract)
    document.add_heading(labels["findings"], level=1)
    for item in findings:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading(labels["taxonomy"], level=1)
    for item in category_lines:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading(labels["top"], level=1)
    for row in top_cited:
        document.add_paragraph(f"{row['year']} #{row['rank']}: {row['title']} ({row['citationCount']:,} citations)", style="List Number")
    document.add_heading(labels["future"], level=1)
    for item in future:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading(labels["refs"], level=1)
    for ref in refs:
        document.add_paragraph(ref, style="List Number")
    document.save(PAPER_DIR / "review_en.docx")


def write_curation_method():
    markdown = f"""# Awesome Brain Curation Method

Generated: {GENERATED_DATE}

## Scope

- Topic: brain research
- Years: {START_YEAR}-{END_YEAR}
- Candidate target: up to {CANDIDATES_PER_YEAR:,} papers per year
- Selection target: {TARGET_PER_YEAR} papers per year
- Ranking: citation count descending, using Semantic Scholar `citationCount`
- Metadata source: Semantic Scholar Academic Graph bulk search, free public metadata

## Query

For each year, the pipeline sends broad brain and neuroscience queries to Semantic Scholar, sorts by citation count, and paginates within the free public endpoint limits. Local relevance checks then keep records whose title, abstract, fields, or publication metadata indicate brain/neuroscience content.

For early years where public metadata contains fewer than the requested target, the generated datasets keep the full audited pool and select all available citation-ranked records rather than fabricating missing entries.

## Enrichment

The script deterministically assigns taxonomy categories, keyword convention tags, key ideas, strengths, and research-focused limitations. No paid API, paid LLM, paid translation service, or paid compute is used.

## GitHub-Awesome Skill2 and Paper-Curation Provenance

This regeneration follows `github-awesome-skill2` in metadata-adapter mode for a large citation-ranked awesome repository while preserving the selected paper set in `data/{PAPERS_CSV}`. The workflow inspected the local `jehyunlee/paper-curation` checkout and is configured for Zotero-free folder-source PDF staging under `E:\\조선대\\연구\\paper-curation\\paper\\awesome-brain`. Full PDF LLM review stages from paper-curation were not run because they require explicit approval for paid or metered APIs.

## Verification Targets

The repository should contain selected and candidate CSV/JSON data, `README.md`, `README.html`, `docs/index.html`, period analysis JSON, taxonomy SVG assets, and English/Korean review HTML files.
"""
    (PAPER_DIR / "curation_method.md").write_text(markdown, encoding="utf-8")
    (PAPER_DIR / "curation_method.html").write_text(markdown_to_html_doc("Awesome Brain Curation Method", markdown), encoding="utf-8")


def write_skill2_provenance(selected, candidates):
    folder_source_pdf_dir = Path(r"E:\조선대\연구\paper-curation\paper\awesome-brain")
    manifest_path = folder_source_pdf_dir / "_folder_source_manifest.json"
    failures_path = folder_source_pdf_dir / "_folder_source_failures.json"
    manifest_count = 0
    failure_count = 0
    if manifest_path.exists():
        manifest_count = len(json.loads(manifest_path.read_text(encoding="utf-8")))
    if failures_path.exists():
        failure_count = len(json.loads(failures_path.read_text(encoding="utf-8")))
    payload = {
        "skill": "github-awesome-skill2",
        "mode": "metadata-adapter",
        "paper_curation_source": "E:\\조선대\\연구\\paper-curation",
        "zotero_used": False,
        "paid_or_metered_api_used": False,
        "folder_source_pdf_dir": str(folder_source_pdf_dir),
        "folder_source_manifest": str(manifest_path),
        "folder_source_manifest_pdfs": manifest_count,
        "folder_source_failed_records": failure_count,
        "selected_dataset": f"data/{PAPERS_CSV}",
        "candidate_dataset": f"data/{CANDIDATES_CSV}",
        "selected_papers": len(selected),
        "candidate_records": len(candidates),
        "period": YEAR_RANGE_TEXT,
        "candidate_target_per_year": CANDIDATES_PER_YEAR,
        "selection_target_per_year": TARGET_PER_YEAR,
        "ranking": "citationCount descending with influentialCitationCount and metadata importance score retained as audit signals",
        "note": "The repository/site outputs are deterministic metadata curation artifacts; full PDF LLM reviews require separate explicit approval.",
    }
    (DATA_DIR / SKILL2_PROVENANCE_JSON).write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def write_citation_and_license():
    citation = f"""cff-version: 1.2.0
title: Awesome Brain
message: If you use this metadata curation, please cite the repository and the original papers.
type: dataset
authors:
  - family-names: Hong
    given-names: Gi
repository-code: https://github.com/honggi82/awesome-brain
url: https://honggi82.github.io/awesome-brain/
date-released: "{GENERATED_DATE}"
license: CC-BY-4.0
"""
    (ROOT / "CITATION.cff").write_text(citation, encoding="utf-8")
    (ROOT / "LICENSE").write_text("CC-BY-4.0 for text and metadata curation; upstream paper metadata belongs to original sources.\n", encoding="utf-8")
    (ROOT / ".gitignore").write_text("__pycache__/\n*.pyc\n.DS_Store\ndata/cache/\n", encoding="utf-8")
    publish = """@echo off
echo This repository is generated locally. Create https://github.com/honggi82/awesome-brain, then run:
echo git remote add origin https://github.com/honggi82/awesome-brain.git
echo git push -u origin main
"""
    (ROOT / "publish_to_github.bat").write_text(publish, encoding="utf-8")


def copy_docs_data():
    for filename in [PAPERS_CSV, PAPERS_JSON, CANDIDATES_CSV, CANDIDATES_JSON, TAXONOMY_CSV, PERIOD_ANALYSIS_JSON, OVERALL_ANALYSIS_JSON, SKILL2_PROVENANCE_JSON]:
        shutil.copyfile(DATA_DIR / filename, DOCS_DIR / "data" / filename)
    if (DATA_DIR / GITHUB_LINKS_JSON).exists():
        shutil.copyfile(DATA_DIR / GITHUB_LINKS_JSON, DOCS_DIR / "data" / GITHUB_LINKS_JSON)
    if (DATA_DIR / LINK_AUDIT_JSON).exists():
        shutil.copyfile(DATA_DIR / LINK_AUDIT_JSON, DOCS_DIR / "data" / LINK_AUDIT_JSON)
    shutil.copyfile(PAPER_DIR / "review_en.html", DOCS_DIR / "paper" / "review_en.html")
    shutil.copyfile(PAPER_DIR / "review_ko.html", DOCS_DIR / "paper" / "review_ko.html")
    shutil.copyfile(PAPER_DIR / "curation_method.html", DOCS_DIR / "paper" / "curation_method.html")


def main():
    refresh = "--refresh" in __import__("sys").argv
    ensure_dirs()
    selected, candidates = collect_papers(refresh=refresh)
    write_json_csv(selected, candidates)
    write_taxonomy_dataset(selected)
    write_period_analysis(selected)
    write_overall_analysis(selected)
    write_taxonomy_assets()
    write_readme(selected, candidates)
    write_site(selected)
    write_review_html(selected, korean=False)
    write_review_html(selected, korean=True)
    write_review_docx(selected)
    write_curation_method()
    write_skill2_provenance(selected, candidates)
    write_citation_and_license()
    copy_docs_data()
    print(f"[done] generated {len(selected):,} selected papers from {len(candidates):,} candidates", flush=True)


if __name__ == "__main__":
    main()
